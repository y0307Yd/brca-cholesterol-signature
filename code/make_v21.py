# -*- coding: utf-8 -*-
"""Generate content-complete v21 from v20.

v21 restores the full methodological and descriptive detail present in v15
(which had been compressed for a word-limited journal), while retaining all
v19/v20 additions: immune deconvolution (CIBERSORT/ESTIMATE), immune
checkpoint genes, LASSO bootstrap stability, GEO subtype external validation,
GSE20711, LIPA/FAXDC2/SREBF1 colocalisation, HPA protein evidence, pan-cancer
FDPS, comparison table, TRIPOD-AI checklist and graphical abstract.

No word-count cap is applied in this version.
"""
import re

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = (r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature"
       r"\Manuscript_Cholesterol_Metabolism_BRCA_v20.docx")
OUT = (r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature"
       r"\Manuscript_Cholesterol_Metabolism_BRCA_v21.docx")


def set_text(p, bold_part, body_part):
    """Replace paragraph runs, preserving paragraph style."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if bold_part:
        r = p.add_run(bold_part.rstrip("\n"))
        r.bold = True
        if bold_part.endswith("\n"):
            p.add_run().add_break()
    lines = body_part.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        if line:
            p.add_run(line)


# key: unique prefix of the v20 paragraph -> (bold part, full merged text)
REPL = {}

REPL["Cholesterol and its derivatives support"] = (
    "",
    "Cholesterol and its derivatives support breast cancer cell "
    "proliferation, membrane biogenesis, steroid hormone synthesis and "
    "immune regulation [1,2]. ER-positive (ER+) tumours coordinately "
    "up-regulate cholesterol biosynthetic enzymes, and cholesterol-lowering "
    "agents have been proposed as adjunct therapy [3-6]. Most studies, "
    "however, examine individual genes and rarely connect pathway-level "
    "expression to patient phenotypes in a structured, multi-layer fashion "
    "[7,8]. Recent lipid-metabolism transcriptomic studies in breast cancer "
    "have emphasised prognostic signatures and immune correlates [7,8]; the "
    "present study differs by reporting an honest null network-pathway "
    "intersection, an ER-focused classifier validated across four "
    "independent expression cohorts, and a germline causal-inference layer "
    "for FDPS.")

REPL["We applied a five-layer bioinformatics framework"] = (
    "",
    "A widely used bioinformatics template proceeds from whole-transcriptome "
    "screening (WGCNA plus a pathway gene set) to machine-learning target "
    "reduction, clinical model building, molecular subtyping, single-cell "
    "localization and Mendelian-randomization causal testing. We applied "
    "this template to breast cancer using cholesterol metabolism as the "
    "biological anchor, with TCGA-BRCA and METABRIC as discovery and "
    "external cohorts, and we report the completed evidence chain from "
    "screening to causal inference.")

REPL["2.1 Data sources and verification"] = (
    "2.1 Data sources and verification ",
    "The overall study design and the flow of analyses are summarised in "
    "Supplementary Figure S2. For validation analyses, PAM50 calls for "
    "TCGA-BRCA were obtained from UCSC Xena (GDC hub) [9]; FDPS HM450 "
    "methylation beta values were obtained from the cBioPortal data "
    "repository (brca_tcga_methylation_hm450); the independent expression "
    "cohorts GSE21653 (Affymetrix HG-U133 Plus 2.0, NCBI GEO with GPL570 "
    "annotation) [10] and GSE7390 (Affymetrix HG-U133A, ArrayExpress "
    "E-GEOD-7390 processed matrix and SDRF clinical data, probes mapped "
    "with hgu133a.db) [11] were used; a fourth independent Affymetrix "
    "cohort (GSE20711, NCBI GEO) was used for classifier transfer; and the "
    "second single-cell atlas GSE161529 (Mendeley Data mirror of the Pal "
    "et al. atlas) [12] was used for single-cell validation. The "
    "Affymetrix annotation was taken from the GEO GPL570 platform file. "
    "TCGA-BRCA RNA-sequencing STAR counts, methylation, mutation and "
    "clinical files were obtained from the NCI GDC portal [13] and verified "
    "against the official manifest by MD5 checksums (4,393/4,393 files); "
    "progression-free interval (PFI) was taken from TCGA-CDR [14]. "
    "METABRIC expression and clinical annotation with recurrence-free "
    "survival (RFS) were obtained from cBioPortal [15]. TCGA ER/PR/HER2 "
    "status was parsed from clinical XML; METABRIC receptor status from "
    "clinical sample files.")

REPL["2.4 WGCNA"] = (
    "2.4 WGCNA ",
    "Expression was transformed as log2(CPM+1). The 8,000 most variable "
    "genes (median absolute deviation) passing expression filters entered "
    "weighted correlation network analysis [21,22]: soft-threshold power "
    "selection by scale-free fit, adjacency |cor|^beta, topological overlap "
    "matrix, average-linkage hierarchical clustering, module cutting with a "
    "height scan (10-40 modules of >= 20 genes) and merging of modules with "
    "eigengene correlation >= 0.80. Module eigengenes were correlated with "
    "PFI event, PFI time, ER status and stage (Pearson/Spearman; "
    "Benjamini-Hochberg FDR [23]).")

REPL["2.5 Candidate selection and machine learning"] = (
    "2.5 Candidate selection and machine learning ",
    "Model calibration was evaluated with reliability curves (decile "
    "observed-versus-predicted ER+ fractions), the Brier score, and "
    "logistic recalibration intercept and slope; clinical utility was "
    "evaluated with decision-curve analysis (net benefit across threshold "
    "probabilities) in TCGA out-of-fold cross-validated predictions and in "
    "the external METABRIC cohort. The strict intersection of "
    "trait-associated modules with the cholesterol set was empty; the "
    "candidate pool was therefore the 119 cholesterol genes measurable in "
    "both cohorts (pathway-first adaptation, documented as an analysis "
    "deviation). Expression was z-scored with TCGA statistics. LASSO "
    "logistic regression [24] (L1, C selected by 5-fold CV over a 25-point "
    "grid) and SVM-RFE (linear SVM) [25] were applied to ER status; the "
    "intersection of their selected gene lists defined 11 hub genes. "
    "Selection stability was assessed with 300 bootstrap resamples of the "
    "LASSO fit. A signature score was computed as the sum of LASSO "
    "coefficients times z-scored expression. Performance was assessed by "
    "5-fold CV AUC in TCGA and external AUC in METABRIC (2,000-bootstrap "
    "95% confidence intervals); as a secondary endpoint, the association "
    "of the signature with PFI/RFS was tested by univariate Cox regression.")

REPL["Consensus clustering [26]"] = (
    "",
    "Consensus clustering [26] (100 resamples, 80% of patients, k-means, "
    "k = 2-6) was applied to hub-gene z-scores; the number of clusters was "
    "selected by the proportion of ambiguous clustering (PAC). Subtypes "
    "were compared for ER composition, PFI (multivariate and pairwise "
    "log-rank), immune marker scores (mean z of curated panels for CD8 T, "
    "CD4 T, regulatory T, NK, B, macrophage, dendritic, neutrophil and "
    "stromal cells), the cholesterol-pathway score, per-gene differential "
    "expression (Mann-Whitney, BH-FDR) and KEGG/GO-BP hypergeometric "
    "enrichment of genes with |mean log2 difference| >= 1. Cross-cohort "
    "validation mapped the TCGA subtype centroids onto METABRIC, GSE21653 "
    "and GSE7390 by nearest-centroid assignment using within-cohort "
    "z-scores of the same hub genes, and compared ER composition, survival "
    "(global and pairwise log-rank) and ER/stage-adjusted subtype "
    "associations (Cox regression, C1 as reference). Immune programs were "
    "additionally quantified by single-sample GSEA (ssGSEA) [27], CIBERSORT "
    "(nu-SVR, LM22, 22 immune cell types) [28] and ESTIMATE (immune/stromal "
    "scores) [29]; immune-checkpoint genes (CD274, PDCD1, CTLA4, LAG3, "
    "HAVCR2, IDO1, CD8A) were compared across subtypes. Hub-gene and FDPS "
    "expression were compared across PAM50 intrinsic subtypes "
    "(Kruskal-Wallis and Mann-Whitney tests).")

REPL["2.8 Single-cell localization and SMR causal inference"] = (
    "2.8 Single-cell localization and SMR causal inference\n",
    "For single-cell localization, we used the processed breast cancer "
    "single-cell atlas of GSE176078 (Wu et al. 2021 [32]; 100,064 cells; "
    "CELLxGENE release) with author-provided cell-type annotations at "
    "three resolutions (9 major, 29 minor and 58 subset types). For each "
    "hub gene we computed mean log-normalized expression, detection "
    "percentage and Wilcoxon enrichment versus all other cells with "
    "Benjamini-Hochberg FDR correction, and confirmed the patterns in "
    "GSE161529 [12]. For causal inference, we used SMR/HEIDI v1.3.1 "
    "[33,34] with eQTLGen whole-blood cis-eQTLs [35] (31,684 individuals; "
    "SMR binary format), the BCAC overall and triple-negative breast "
    "cancer GWAS [36] (247,173 and 118,987 women) and the 1000 Genomes "
    "Phase 3 European LD reference, scanning 119 cholesterol genes (251 "
    "probes) with a +/-2 Mb cis window, instrument eQTL P < 5e-8, MAF > "
    "0.01, HEIDI method 1 and Bonferroni alpha = 0.05/119 = 0.00042. "
    "ER-stratified sensitivity used the BCAC OncoArray public release "
    "[37] (ER+ 69,501; ER- 21,468 cases). GTEx v8 breast-mammary and "
    "whole-blood SMR-format BESDs [38,39] were used for tissue-specific "
    "and replication analyses; expected replication power was calculated "
    "from the eQTLGen instrument under sample-size scaling "
    "(z_GTEx = z_eQTLGen x sqrt(N_GTEx/N_eQTLGen)). For the FDPS locus, "
    "Bayesian colocalisation (coloc.abf; p1 = p2 = 1e-4, p12 = 1e-5, "
    "W = 0.2) was applied across the +/-2 Mb cis window and in +/-1 Mb, "
    "+/-500 kb and +/-250 kb sensitivity windows, with per-SNP posterior "
    "probabilities of a shared causal variant (PP.H4) calculated as in the "
    "published coloc.abf model; single-SNP conditional analysis used the "
    "1000 Genomes EUR reference panel (conditional z-scores computed as "
    "(z - r*z_cond)/sqrt(1 - r^2) with signed genotype correlations). "
    "Multi-signal sensitivity used SuSiE fine-mapping [40] on the eQTLGen "
    "and BCAC z-scores with 1000 Genomes EUR LD, followed by coloc.susie "
    "[41]. Detailed parameter settings are provided in the archived code.")

REPL["At soft-threshold power"] = (
    "",
    "At soft-threshold power beta = 5 (scale-free fit R2 = 0.834, "
    "slope = -1.64), the 8,000-gene network resolved into 15 modules "
    "(38-853 genes). The module most strongly associated with PFI events "
    "(M2238; 69 genes) showed a weak positive correlation (r = 0.082, "
    "P = 0.012) that did not survive correction for multiple testing "
    "(BH-FDR = 0.170); modules M3105 (r = -0.071) and M2722 (r = 0.069) "
    "were comparable. In contrast, ER status showed very strong module "
    "associations, most notably M2314 (157 genes; r = -0.794, "
    "BH-FDR ~ 4e-197).")

REPL["The strict intersection of trait-associated modules with the 120-gene"] = (
    "",
    "Strictly intersecting any trait-associated module with the 120-gene "
    "cholesterol set yielded zero genes: the PFI module (M2238) and the ER "
    "module (M2314) contained no cholesterol metabolism genes, and only 15 "
    "of 55 cholesterol genes present in the network fell into formal "
    "modules (the remainder resided in small clusters excluded from module "
    "calling). At the single-gene level, no cholesterol gene was "
    "significantly associated with PFI (smallest Cox P = 0.13, CH25H), "
    "whereas several strongly tracked ER status (LIPE P = 4.5e-17; EBP "
    "P = 2.4e-12; FDXR P = 6.0e-12; PMVK P = 1.4e-05; SQLE P = 1.5e-04). "
    "This is reported as an honest null result for the strict template and "
    "motivated the pathway-first adaptation below.")

REPL["3.3 Hub genes and the ER-status classifier"] = (
    "3.3 Hub genes and the ER-status classifier ",
    "From 119 measurable cholesterol genes, LASSO (C = 0.316; 69 genes) "
    "and SVM-RFE (12 genes) selected 11 hub genes by intersection "
    "(Table 3). In 300 LASSO bootstrap resamples all 11 hub genes were "
    "selected in at least 73.7% of fits (mean 90.2%; G6PD and VLDLR 100%, "
    "PRKAA1 99.7%, LIMA1 73.7%), confirming selection stability. The "
    "ER-status classifier reached internal 5-fold CV AUC 0.946 +/- 0.012 "
    "(full-data 0.953) and external AUC 0.922 (95% CI 0.907-0.935) in "
    "METABRIC. Calibration was excellent in TCGA out-of-fold "
    "cross-validation (Brier 0.069; recalibration intercept 0.002, slope "
    "1.06). In METABRIC, discrimination remained high but external "
    "calibration showed mild systematic underestimation of ER+ probability "
    "(Brier 0.098; intercept 0.57, slope 1.41), indicating that "
    "cohort-specific recalibration would be appropriate before clinical "
    "use; decision-curve analysis showed positive net benefit across "
    "threshold probabilities of 0.05-0.95 in both cohorts (Figure 3). The "
    "signature was not associated with PFI in TCGA (HR 0.89 per SD, 95% CI "
    "0.78-1.01, P = 0.071; adjusted for ER, stage and age: HR 0.90, "
    "P = 0.151) but showed a protective association with RFS in METABRIC "
    "(HR 0.88 per SD, 95% CI 0.83-0.93, P = 4.2e-05), which persisted "
    "after adjustment for ER status and stage (HR 0.90, 95% CI 0.82-0.98, "
    "P = 0.013).")

REPL["Consensus clustering selected k = 4"] = (
    "",
    "Consensus clustering of hub-gene expression selected k = 4 "
    "(PAC = 0.097; k = 3 PAC 0.113, k = 5 PAC 0.185). The four subtypes "
    "differed markedly in ER composition (C1 13.6%, C2 93.8%, C3 97.5%, "
    "C4 60.1% ER+), signature score and cholesterol-pathway score "
    "(Table 4). Survival differences were not significant in TCGA "
    "(multivariate log-rank P = 0.331; pairwise BH-FDR >= 0.39). "
    "Nearest-centroid mapping of the TCGA subtype centroids onto METABRIC "
    "(within-cohort hub-gene z-scores) reproduced all four subtypes with "
    "near-identical ER composition (C1 15.6%, C2 92.6%, C3 94.0%, C4 50.1% "
    "ER+ versus 13.6%, 93.8%, 97.5% and 60.1% in TCGA) and with "
    "significant survival differences across subtypes (global log-rank "
    "P = 7.8e-05; C4 worst, C2 best; C2 vs C4 P = 4.1e-06). After "
    "adjustment for ER status and stage in multivariable Cox analysis, no "
    "subtype differed significantly from C1 (C4 vs C1 HR = 1.26, 95% CI "
    "0.95-1.67, P = 0.12), indicating that the unadjusted survival "
    "differences largely tracked ER status and stage (Figure 7). "
    "Nearest-centroid mapping onto GSE21653 (n = 252) and GSE7390 "
    "(n = 198) reproduced the subtype ER gradient (Kruskal-Wallis "
    "P = 4.2e-16 and P = 4.7e-16, respectively); survival differences were "
    "significant in GSE21653 (global log-rank P = 0.026; C1 worst, C2 "
    "best) but not in GSE7390 (P = 0.48), consistent with ER-driven "
    "survival differences (Supplementary Figure S3).")

REPL["3.5 Immune microenvironment and pathway programs"] = (
    "3.5 Immune microenvironment and pathway programs across subtypes\n",
    "All nine immune/stromal marker scores differed significantly across "
    "subtypes (Kruskal-Wallis P < 0.001): C1 showed the highest CD8 T, "
    "regulatory T, NK, B-cell and macrophage scores and the lowest stromal "
    "score; C3 the lowest across most immune populations; C2 the highest "
    "dendritic-cell and stromal scores; C4 the highest neutrophil score "
    "(Table 5). The same scores computed in METABRIC reproduced the TCGA "
    "pattern in the mapped subtypes (all Kruskal-Wallis P < 0.001; "
    "Supplementary Figure S1). Differential expression (each subtype vs "
    "the rest, BH-FDR < 0.05) identified 28,931 / 28,287 / 23,884 / "
    "15,975 genes for C1-C4, reflecting the strong ER-driven transcriptome "
    "shift; restricting to |mean log2 difference| >= 1 gave 1,840 / 304 / "
    "225 / 92 genes. Enrichment of the strong-effect genes showed that C1 "
    "was characterized by up-regulated cell-cycle programs (GO: nuclear "
    "division, chromosome segregation, mitotic cell cycle phase "
    "transition; KEGG hsa04110 Cell cycle), C2 by down-regulated "
    "chromosome-segregation/mitotic programs, C3 by up-regulated oxidative "
    "phosphorylation and ribosomal programs, and C4 by up-regulated "
    "protein-processing, cell-cycle and proteasome programs. ssGSEA "
    "enrichment scores recapitulated these patterns and confirmed "
    "significant subtype differences for CD8 T, CD4 T, regulatory T, NK, "
    "macrophage, dendritic-cell, neutrophil and stromal programs (all "
    "Kruskal-Wallis P < 0.05), with B-cell scores as the exception "
    "(P = 0.42); correlations between ssGSEA and marker scores ranged from "
    "0.40 to 0.93 (Figure 8; Supplementary Table S8). ESTIMATE confirmed "
    "the immune-hot/immune-cold axis (ImmuneScore highest in C1, P = "
    "1.9e-3; StromalScore highest in C2, P = 4.2e-22; ESTIMATEScore "
    "P = 3.5e-6). CIBERSORT showed higher C1 fractions of M1/M0 "
    "macrophages, activated dendritic cells, follicular helper T cells, "
    "monocytes and activated NK cells, and higher C2/C3 fractions of "
    "resting mast cells and M2 macrophages (all P < 0.001). "
    "Immune-checkpoint genes CD274, PDCD1, CTLA4, LAG3 and IDO1 were "
    "highest in C1 (P = 1.7e-20 to 4.9e-5); HAVCR2 and CD8A did not "
    "differ (P = 0.12 and 0.11; Supplementary Tables S8 and S14).")

REPL["3.6 Single-cell localization of hub genes"] = (
    "3.6 Single-cell localization of hub genes\n",
    "In the 100,064-cell breast cancer atlas, all 11 hub genes were "
    "detected. Cholesterol biosynthetic enzymes were most strongly enriched "
    "in malignant epithelial cells: DHCR24 (mean 0.45, 46.9% of cells), "
    "G6PD (0.27, 32.0%), DHCR7 (0.26, 31.8%), HSD17B7 (0.18, 25.2%) and "
    "FDXR (0.12, 18.6%), all with Wilcoxon BH-FDR < 1e-10. ABCG1 was "
    "enriched in myeloid cells (mean 0.24, 25.9%), specifically "
    "macrophages (0.30, 31.8%), and in endothelial cells (0.26, 25.9%); "
    "G6PD (0.18, 22.2%) and LIMA1 (0.12, 15.2%) were also enriched in "
    "myeloid cells, and LIMA1 in cancer-associated fibroblasts (0.75, "
    "52.2%). These patterns localize cholesterol metabolism to both the "
    "tumour epithelial compartment and the immune/stromal microenvironment.")

REPL["3.7 SMR causal inference"] = (
    "3.7 SMR causal inference\n",
    "A pathway-wide SMR scan of 119 cholesterol genes identified three "
    "probes passing the Bonferroni threshold in overall breast cancer: "
    "CYP51A1 (P = 1.4e-8), FDPS (P = 9.8e-7) and LSS (P = 4.1e-5); "
    "SREBF1 (P = 4.6e-4) was borderline. HEIDI rejected the CYP51A1 "
    "(P = 6.5e-6) and LSS (P = 4.2e-5) signals as LD-confounded, whereas "
    "FDPS passed (P_HEIDI = 0.026), with higher genetically predicted "
    "FDPS expression associated with lower breast cancer risk (b = -0.42, "
    "95% CI -0.59 to -0.25). SREBF1 (P_HEIDI = 0.17), LIPA (P = 5.8e-3; "
    "P_HEIDI = 0.77) and FAXDC2 (P = 3.3e-3; P_HEIDI = 0.13) showed the "
    "same protective direction with HEIDI support. In the TNBC/BRCA1 "
    "sensitivity analysis, LSS was nominally significant (P = 1.4e-3) "
    "with the same direction; no probe passed Bonferroni correction. "
    "G6PD, HMGCS2 and NSDHL lacked eQTLGen instruments. Full per-probe "
    "SMR and HEIDI results for the 81 tested probes across eQTL resources "
    "(eQTLGen whole blood, GTEx v8 breast-mammary and GTEx v8 whole "
    "blood) are provided in Supplementary Table S2. In the ER-stratified "
    "OncoArray sensitivity analysis, the FDPS protective direction was "
    "replicated: ER+ P = 1.54e-05 (b = -0.44, HEIDI P = 1.09e-02, top SNP "
    "rs6677385) and ER- P = 2.78e-01 (b = -0.15, HEIDI P = 6.54e-01, top "
    "SNP rs6677385); ER+ reached Bonferroni significance, whereas ER- was "
    "underpowered (Supplementary Table S10). In the OncoArray overall "
    "meta-analysis, FDPS showed the same protective direction "
    "(P = 9.02e-06, HEIDI P = 3.98e-02).")

REPL["3.8 GTEx breast-tissue SMR sensitivity"] = (
    "3.8 GTEx breast-tissue SMR sensitivity analysis\n",
    "To test tissue specificity, we repeated the SMR and HEIDI analyses "
    "with SMR-format cis-eQTL summary data from GTEx v8 breast-mammary "
    "tissue (24,290 probes; GRCh37; Yang Lab SMR database), the same BCAC "
    "GWAS (overall and TNBC) and the same 1000 Genomes EUR LD reference, "
    "restricting GWAS variants to the cis regions of the 132 "
    "cholesterol-gene probes (+/-2 Mb). Of the 119 cholesterol genes with "
    "eQTLGen instruments, 109 had breast-tissue eQTL probes and 16 had "
    "genome-wide-significant cis-eQTLs (P < 5e-8) available for testing. "
    "No probe passed Bonferroni correction in either analysis (threshold "
    "P = 3.1e-3). Two probes replicated direction-consistently at nominal "
    "significance in overall breast cancer with HEIDI support: LIPA "
    "(P = 8.2e-3; P_HEIDI = 0.53) and FAXDC2 (P = 1.15e-2; P_HEIDI = "
    "0.69), both matching the eQTLGen blood results (LIPA P = 5.8e-3, "
    "P_HEIDI = 0.77; FAXDC2 P = 3.3e-3, P_HEIDI = 0.13) with the same "
    "positive SMR direction. In the TNBC analysis, PON1 was nominally "
    "significant (P = 4.2e-2; P_HEIDI = 0.49). FDPS, CYP51A1, LSS and "
    "SREBF1 lacked genome-wide-significant breast-tissue cis-eQTL "
    "instruments, indicating that the blood-based FDPS signal is likely "
    "mediated through non-mammary (e.g., immune or metabolic) tissues, "
    "consistent with tissue-specific regulation of cholesterol metabolism.")

REPL["3.9 Colocalisation supports"] = (
    "3.9 Colocalisation supports a shared causal variant at the FDPS locus\n",
    "To distinguish a shared causal variant from coincident but "
    "independent signals, we applied Bayesian colocalisation (coloc.abf "
    "model; p1 = p2 = 1e-4, p12 = 1e-5, W = 0.2) to the FDPS cis region "
    "(chr1q22; probe at 155,284,498 bp; +/-2 Mb). Of 4,215 cis-eQTL SNPs "
    "and 7,880 GWAS SNPs in the region, 3,248 were shared and "
    "informative. The posterior probability of a shared causal variant "
    "(PP.H4) was 0.9985; narrowing the window to +/-1 Mb, +/-500 kb and "
    "+/-250 kb gave PP.H4 = 0.997, 0.996 and 0.997, respectively. The lead "
    "SNP rs12091730 (chr1:155,556,971) was strongly associated with both "
    "FDPS blood expression (P = 2.5e-18) and breast cancer risk "
    "(P = 1.8e-12) with opposing effect directions, consistent with the "
    "protective SMR estimate (higher genetically predicted FDPS "
    "expression, lower risk). The strongest shared signals formed a "
    "single LD block of 11 SNPs with r2 >= 0.8 with rs12091730 "
    "(chr1:155.41-155.67 Mb; 1000 Genomes EUR), and the correlation "
    "between eQTL and GWAS Z-statistics across the region was -0.65. A "
    "secondary GWAS-only signal (rs4971059, chr1:155,148,781; "
    "P = 2.5e-12) showed only a weak eQTL effect (P = 9.7e-4), and a "
    "further GWAS signal (rs11264454, chr1:156,153,043) remained "
    "genome-wide significant after conditioning on the shared lead "
    "(P = 2.7e-10) with no eQTL effect (P = 0.55). Single-SNP conditional "
    "analysis with 1000 Genomes EUR LD showed that the FDPS cis-eQTL is a "
    "single signal: after conditioning on rs12091730, no SNP remained "
    "genome-wide significant (minimum conditional P = 2.1e-06). "
    "Conversely, conditioning on the secondary GWAS signal rs4971059 left "
    "the shared lead genome-wide significant (conditional P = 4.4e-08), "
    "and conditioning on rs12091730 left rs4971059 and rs11264454 "
    "significant (conditional P = 6.0e-08 and 2.7e-10). The shared "
    "eQTL-GWAS signal is therefore robust to the additional GWAS "
    "variants, which lack a blood-eQTL effect (Supplementary Table S1). "
    "Sensitivity analyses with the alternative prior p12 = 1e-8 gave "
    "PP.H4 = 0.40 (PP.H3 = 0.38), whereas all settings at the recommended "
    "prior p12 = 1e-5 gave PP.H4 >= 0.99 (Figure 9). The FDPS eQTL "
    "instrument was strong (top cis-eQTL SNP rs6677385; F = 93.1), and "
    "the 99% credible set for the shared eQTL-GWAS signal comprised 56 "
    "SNPs within the 155.28-155.56 Mb LD block (cumulative SNP-level "
    "PP.H4 = 0.99; Supplementary Table S1). However, multi-signal "
    "SuSiE/coloc.susie identified four eQTL credible sets and one GWAS "
    "credible set with maximum PP.H4 = 0.44 (PP.H3 = 0.56), so the "
    "single-shared-variant interpretation of coloc.abf should be regarded "
    "as suggestive (Supplementary Table S1). Bayesian colocalisation was "
    "extended to LIPA, FAXDC2 and SREBF1; none supported a shared causal "
    "variant with breast cancer risk (PP.H4 = 0.035, 0.043 and 0.171; "
    "PP.H1 dominated), indicating strong eQTL but no corresponding GWAS "
    "signal (Supplementary Table S15).")

REPL["3.10 FDPS expression and clinical features"] = (
    "3.10 FDPS expression and clinical features\n",
    "FDPS expression was strongly associated with ER status in both "
    "cohorts: ER-negative tumours expressed FDPS at higher levels than "
    "ER-positive tumours (TCGA-BRCA P = 5.7e-17, Cohen's d = -0.69; "
    "METABRIC P = 2.2e-48, d = -0.79), and FDPS was not associated with "
    "tumour stage (TCGA rho = 0.04, P = 0.19). In univariable Cox "
    "analysis higher FDPS expression was associated with worse "
    "recurrence-free survival in METABRIC (HR = 1.09 per SD, P = 0.015) "
    "but not in TCGA-BRCA (P = 0.79); after adjustment for ER status and "
    "stage the METABRIC association was attenuated to the null (HR = "
    "1.03, P = 0.48). FDPS expression therefore tracks ER status rather "
    "than acting as an independent tumour prognostic factor, and the "
    "germline eQTL-based causal signal is distinct from this tumour-level "
    "association. In the Human Protein Atlas [42], FDPS showed cytoplasmic "
    "protein localization, tissue-enhanced expression in normal breast, "
    "and a non-significant TCGA protein-level prognosis (P = 0.195; HPA "
    "validation P = 0.023, classified unprognostic), consistent with the "
    "transcript-level null. FDPS mRNA was detected in all 17 queried "
    "cancer types (Supplementary Table S13).")

REPL["3.11 GTEx v8 whole-blood sensitivity"] = (
    "3.11 GTEx v8 whole-blood sensitivity analysis\n",
    "We attempted to replicate the FDPS instrument in an independent "
    "whole-blood resource, GTEx v8 (n = 670 [38]; SMR-format BESD from "
    "the Yang Lab SMR database, 19,270 probes). No cis-eQTL records were "
    "available for the FDPS probe in this BESD, and the GTEx Portal v2 "
    "single-tissue eQTL endpoint [39] returned no significant FDPS "
    "cis-eQTL in whole blood (FDR < 0.05), confirming that FDPS is not an "
    "eGene in this smaller resource. This is expected under sample-size "
    "scaling: at the eQTLGen instrument (z = -9.65, N = 31,684), the "
    "expected GTEx whole-blood z-statistic is -1.40, giving 29% power at "
    "P < 0.05 and <0.01% power at P < 5e-8 (rs12091730: expected z = "
    "-1.27, 25% power at P < 0.05). The FDPS causal estimate therefore "
    "relies on the larger eQTLGen whole-blood resource, and GTEx v8 whole "
    "blood cannot provide an independent instrument for this gene.")

REPL["PAM50 calls were available for 871"] = (
    "",
    "PAM50 calls were available for 871 TCGA-BRCA tumours. The hub-gene "
    "subtypes differed markedly in PAM50 composition (chi-square = 820.8, "
    "P = 5.62e-168, Cramer's V = 0.560): C1 was predominantly Basal "
    "(86.4%), C2 LumA (67.4%) and C3 LumA (48.6%), demonstrating that "
    "the expression-based subtypes map onto the clinically established "
    "intrinsic subtypes (Figure 10; Supplementary Table S4). Hub-gene and "
    "FDPS expression differed across all five PAM50 subtypes "
    "(Kruskal-Wallis P < 0.001 for all 12 genes); FDPS expression was "
    "highest in Basal and HER2-enriched tumours and lowest in Normal-like "
    "tumours, consistent with its ER-negative enrichment (Supplementary "
    "Table S9).")

REPL["3.13 Independent validation in GSE21653"] = (
    "3.13 Independent validation in GSE21653\n",
    "In an independent Affymetrix GPL570 cohort (GSE21653, n = 252 [10], "
    "83 disease-free-survival events), 11/11 hub genes were mapped to "
    "probes and per-gene inverse-normal transformed before applying the "
    "TCGA-fitted classifier. The ER classifier transferred with "
    "AUC = 0.847 (95% CI 0.798-0.895; Brier 0.210), and FDPS expression "
    "was strongly associated with ER status (Mann-Whitney P = 1.24e-10; "
    "Cohen's d = -0.86). In survival analysis, the direction observed in "
    "TCGA/METABRIC was not replicated: higher FDPS expression was "
    "nominally associated with shorter DFS in univariable analysis "
    "(HR per SD = 1.27, 95% CI 1.01-1.58, P = 0.037), but the association "
    "attenuated after adjustment for ER status, age and grade (HR = 1.16, "
    "P = 0.250; median-split log-rank P = 0.295), suggesting that the "
    "crude effect was largely driven by the strong inverse FDPS-ER "
    "relationship in this predominantly hormone-receptor-negative "
    "surgical series. The classifier transfer and ER biology therefore "
    "replicate across platforms, whereas the prognostic direction is "
    "cohort-dependent (Figures 12 and 13; Supplementary Table S5). In a "
    "third independent cohort (GSE7390 [11], n = 198, 91 RFS events), "
    "the classifier transferred with AUC = 0.903 (95% CI 0.848-0.949; "
    "Brier 0.119), and FDPS remained strongly ER-associated "
    "(Mann-Whitney P = 4.78e-05; d = -0.68), confirming cross-platform "
    "reproducibility. In a fourth independent cohort (GSE20711, n = 88 "
    "with RFS and 39 events; Fackler et al. [43]), the ER classifier "
    "transferred with AUC = 0.833 (95% CI 0.739-0.913; Brier 0.226) and "
    "FDPS remained ER-associated (Mann-Whitney P = 2.4e-3; d = -0.64) "
    "with no FDPS survival association (HR = 0.97, P = 0.86), confirming "
    "cross-platform transferability (Supplementary Table S16).")

REPL["3.14 Second single-cell atlas confirms"] = (
    "3.14 Second single-cell atlas confirms FDPS enrichment in "
    "monocyte/macrophage populations\n",
    "In the independent GSE161529 atlas (8 10x profiles; 70,419 cells "
    "after QC), FDPS was most strongly expressed in monocyte/macrophage "
    "populations: Monocyte (mean 0.65, FDR 0.00e+00); Other (mean 0.15, "
    "FDR 0.00e+00); CD8_T (mean 0.29, FDR 2.50e-76). Co-expression of "
    "ligand-receptor pairs linking T cells to myeloid cells was detected "
    "(SPP1-CD44 (7.1% T cells / 0.9% myeloid cells)), consistent with "
    "the active T-cell-macrophage crosstalk observed in the primary "
    "atlas (Supplementary Table S6). An expanded T-myeloid "
    "ligand-receptor analysis of 20 curated pairs identified 7 expressed "
    "pairs; the strongest co-expression was IL1B-IL1R2 (11.2% of "
    "T/myeloid cells), CXCL9/10/11-CXCR3 (2.0-2.3%), CSF1-CSF1R (1.7%) "
    "and IL18-IL18R1 (0.9%) (Figure 14; Supplementary Table S6).")

REPL["3.15 FDPS DNA methylation is inversely associated"] = (
    "3.15 FDPS DNA methylation is inversely associated with FDPS "
    "expression in TCGA-BRCA\n",
    "In 672 TCGA-BRCA tumours with HM450 methylation and RNA-seq data, "
    "FDPS methylation (mean gene-associated beta, cBioPortal) showed a "
    "significant negative correlation with FDPS expression (Spearman "
    "rho = -0.203, 95% CI -0.274 to -0.133, P = 1.04e-07), i.e. higher "
    "methylation tended to accompany lower expression. The effect was "
    "moderate in magnitude (mean beta 0.018) and is consistent with the "
    "expected epigenetic silencing relationship; FDPS methylation "
    "differed by ER status (ER+ 0.018 vs ER- 0.017, Mann-Whitney "
    "P = 0.032), partly mirroring the ER dependence of FDPS expression. "
    "After adjustment for ER status, the inverse methylation-expression "
    "correlation persisted (partial rho = -0.187, P = 2.2e-6) and was "
    "significant in ER+ tumours (rho = -0.206, P = 4.6e-6), with a "
    "consistent but non-significant direction in ER- tumours "
    "(rho = -0.142, P = 0.087; n = 147) (Figures 15 and 16; "
    "Supplementary Table S7).")

REPL["4.1 Principal findings"] = (
    "4.1 Principal findings ",
    "This study applied a five-layer bioinformatics framework to "
    "cholesterol metabolism in breast cancer using two large, verified "
    "cohorts. The strict WGCNA-by-pathway intersection was empty; a "
    "pathway-first adaptation produced an 11-gene cholesterol signature "
    "that discriminated ER status with strong, cross-platform performance "
    "(TCGA CV AUC 0.946 +/- 0.012; METABRIC 0.922; GSE21653 0.847; "
    "GSE7390 0.903; GSE20711 0.833). Hub-gene-based consensus clustering "
    "defined four subtypes with strongly divergent ER composition and "
    "immune infiltration, reproduced in METABRIC, GSE21653 and GSE7390 "
    "and supported by CIBERSORT/ESTIMATE deconvolution and immune-"
    "checkpoint genes. Single-cell analysis localized the signature to "
    "tumour epithelial and myeloid/endothelial compartments, and a "
    "pathway-wide SMR scan identified FDPS as a putatively causal "
    "protective gene for breast cancer risk, with SREBF1/LIPA/FAXDC2 "
    "providing supportive evidence. Bayesian colocalisation supported a "
    "shared causal variant at the FDPS locus (PP.H4 = 0.9985), while "
    "tumour FDPS expression was strongly ER-associated but not "
    "independently prognostic.")

REPL["GTEx v8 breast-mammary sensitivity"] = (
    "",
    "Each layer of the framework was completed with real data. First, the "
    "strict WGCNA-by-pathway intersection was empty: cholesterol "
    "metabolism genes do not participate in the PFI-associated "
    "co-expression module, and no single cholesterol gene was associated "
    "with progression at P < 0.10. Second, a pathway-first adaptation "
    "produced an 11-gene cholesterol signature that discriminated ER "
    "status with strong, cross-platform performance (TCGA CV AUC 0.946 "
    "+/- 0.012; METABRIC AUC 0.922; GSE21653 0.847; GSE7390 0.903; "
    "GSE20711 0.833). Third, hub-gene-based consensus clustering defined "
    "four subtypes with strongly divergent ER composition and immune "
    "infiltration, although survival differences were not significant in "
    "the 122-event discovery cohort. A GTEx v8 breast-mammary sensitivity "
    "analysis directionally replicated LIPA and FAXDC2 (both "
    "HEIDI-supported) but could not instrument FDPS in breast tissue; "
    "GTEx v8 whole blood lacked a significant FDPS cis-eQTL instrument, "
    "consistent with the power available at n = 670.")

REPL["The ER-status signal is biologically coherent"] = (
    "",
    "The ER-status signal is biologically coherent: ER+ tumours "
    "up-regulate cholesterol biosynthetic enzymes to support steroid "
    "production and membrane turnover [3-5], and several hub genes "
    "(DHCR7, DHCR24, NSDHL, HMGCS2, FDXR, HSD17B7) are direct enzymes or "
    "regulators of cholesterol/steroid synthesis. The external AUC of "
    "0.922 in METABRIC - a different platform, cohort and measurement "
    "generation - confirms that the signal is not a TCGA artifact. At the "
    "same time, the near-null PFI associations indicate that cholesterol "
    "gene expression adds little independent prognostic information in "
    "these cohorts, a finding consistent with our earlier "
    "leakage-controlled multi-omics evaluation in which expression "
    "features did not improve prognostic discrimination beyond clinical "
    "variables. FDPS methylation was inversely correlated with expression "
    "independently of ER status, and FDPS expression was highest in Basal "
    "and HER2-enriched intrinsic subtypes (Figure 10).")

REPL["The four subtypes recapitulate known breast cancer biology"] = (
    "",
    "The four cholesterol-metabolism subtypes recapitulate known breast "
    "cancer biology: C1 is predominantly ER-negative, proliferative "
    "(cell-cycle up) and immune-infiltrated (highest CD8 T, regulatory T, "
    "NK, B-cell and macrophage scores), consistent with the "
    "immunologically hot, proliferative phenotype of triple-negative "
    "tumours; C3 is almost uniformly ER+, immune-cold and "
    "oxidative-phosphorylation-high; C2 shows down-regulated cell-cycle "
    "programs and high stromal/dendritic scores; C4 is a mixed, "
    "neutrophil-high group. These patterns provide a cell-level "
    "interpretation of the signature even before single-cell data are "
    "added: cholesterol-synthesis-high subtypes are luminal-like and "
    "immune-cold, whereas cholesterol-synthesis-low subtypes are more "
    "heterogeneous, proliferative and immune-active [2,8]. ssGSEA "
    "confirmed that C1 is immune-hot and C3 immune-cold, providing a "
    "rank-based, non-parametric complement to the marker-score analysis.")

REPL["4.5 Limitations"] = (
    "4.5 Limitations\n",
    "First, the strict WGCNA-by-pathway intersection was empty, and the "
    "candidate pool was therefore defined by the pathway-first "
    "adaptation; this deviation is documented and does not affect the "
    "ER-discrimination conclusion, but it weakens the claim that "
    "WGCNA-guided prioritization was essential. Second, an ER classifier "
    "built from expression partly rediscovers ER-pathway biology; its "
    "novelty is descriptive rather than mechanistic. Third, TCGA had only "
    "122 PFI events, so subtype survival differences were underpowered; "
    "the METABRIC survival association (P = 4.2e-05 univariable; "
    "P = 0.013 after ER/stage adjustment) is the stronger prognostic "
    "evidence, and subtype survival differences were not significant in "
    "GSE7390 (global log-rank P = 0.48). Fourth, immune estimates were "
    "obtained from curated marker panels, ssGSEA, CIBERSORT and ESTIMATE "
    "[28,29]; GO term names were curated offline and should be "
    "re-verified before submission. Fifth, single-cell localization used "
    "two atlases and is descriptive; the cell-type enrichments and "
    "ligand-receptor co-expression are observational, and the HPA protein "
    "evidence is limited to descriptive localization with a "
    "non-significant protein-level prognosis. Sixth, SMR was performed "
    "with blood (eQTLGen) cis-eQTLs; breast-tissue instruments were not "
    "available, and tissue-specific causal effects cannot be excluded. "
    "The strong CYP51A1 and LSS SMR signals were rejected by HEIDI and "
    "should not be interpreted as causal; the FDPS HEIDI P-value (0.026) "
    "was not adjusted for the number of probes tested, and no independent "
    "whole-blood resource currently provides a genome-wide-significant "
    "FDPS cis-eQTL (GTEx v8 whole blood, n = 670, contains no significant "
    "FDPS eQTL). G6PD, HMGCS2 and NSDHL lacked eQTLGen instruments; TNBC "
    "sensitivity analysis found no Bonferroni-significant probe. "
    "Colocalisation with the single-causal-variant coloc.abf model gave "
    "PP.H4 = 0.9985, but SuSiE-based multi-signal colocalisation [40,41] "
    "was less conclusive (maximum PP.H4 = 0.44), so the shared-variant "
    "interpretation should be regarded as suggestive. Conditional "
    "analysis confirmed additional GWAS-only signals at the FDPS locus "
    "(rs4971059 and rs11264454) with no residual eQTL effect after "
    "conditioning on the shared lead. FDPS tumour expression was strongly "
    "ER-associated and not an independent prognostic factor, so the "
    "causal claim is confined to germline-regulated blood expression "
    "rather than tumour expression.")

REPL["4.6 Causal inference"] = (
    "4.6 Causal inference ",
    "By SMR with eQTLGen and BCAC summary statistics, causal inference "
    "was completed for 81 cholesterol-gene probes; FDPS passed Bonferroni "
    "correction and HEIDI, and Bayesian coloc.abf supported a shared "
    "causal variant at the FDPS locus (PP.H4 = 0.9985), together "
    "indicating a protective effect of genetically determined FDPS "
    "expression on breast cancer risk [44-46]. The protective direction "
    "was replicated in ER+ (P = 1.54e-05, HEIDI P = 1.09e-02) and ER- "
    "(P = 2.78e-01) OncoArray sensitivity analyses [37], with ER+ "
    "reaching Bonferroni significance and ER- underpowered. These results "
    "are consistent with Mendelian-randomization evidence that statins "
    "reduce ER+ breast cancer risk [47] and with the established adjuvant "
    "benefit of nitrogen-containing bisphosphonates, which inhibit FDPS "
    "[48]. GTEx whole-blood replication and ER-stratified GWAS remain "
    "possible sensitivity analyses; multi-signal colocalisation was less "
    "conclusive (maximum PP.H4 = 0.44).")


def main():
    doc = docx.Document(SRC)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    matched = set()
    for key, (bold, body) in REPL.items():
        hits = [p for p in paras if p.text.strip().startswith(key)]
        assert len(hits) == 1, f"{key!r}: {len(hits)} matches"
        set_text(hits[0], bold, body)
        matched.add(hits[0]._p)
    print("replaced", len(matched), "paragraphs")
    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()
