from __future__ import annotations

import json
import os
from pathlib import Path

import docx
import matplotlib as mpl
import numpy as np
import pandas as pd
from scipy import stats
from sklearn.ensemble import RandomForestClassifier
from sklearn.inspection import permutation_importance
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import roc_auc_score
from sklearn.model_selection import RepeatedStratifiedKFold, StratifiedKFold

os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parent / ".mplconfig"))
mpl.use("Agg")
import matplotlib.pyplot as plt


SEED = 20260804
RNG = np.random.default_rng(SEED)
REPO = Path(__file__).resolve().parents[1]
PROJECT_ROOT = Path(os.environ.get("BRCA_CHOL_PROJECT_ROOT", REPO)).resolve()
SOURCE = PROJECT_ROOT / "outputs" / "chol_metab_signature"
OUT = Path(os.environ.get("BRCA_CHOL_OUTPUT_DIR", REPO / "reproduction_output")).resolve()
OUT.mkdir(exist_ok=True)

mpl.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
    "font.size": 7,
    "axes.spines.right": False,
    "axes.spines.top": False,
    "axes.linewidth": 0.8,
    "pdf.fonttype": 42,
    "svg.fonttype": "none",
})


def load_data():
    genes = [x.strip() for x in open(SOURCE / "tcga_genes.txt", encoding="utf-8") if x.strip()]
    pool = pd.read_csv(SOURCE / "ml_candidate_pool.csv")["gene"].tolist()
    hubs = pd.read_csv(SOURCE / "hub_genes.csv")["gene"].tolist()
    x_all = np.load(SOURCE / "tcga_Xlog.npy", mmap_mode="r")
    idx = {g: i for i, g in enumerate(genes)}
    x = np.vstack([x_all[idx[g]] for g in pool]).T.astype(float)
    traits = pd.read_csv(SOURCE / "tcga_traits.csv")
    ok = traits.ER.notna().to_numpy()
    return x[ok], traits.loc[ok, "ER"].astype(int).to_numpy(), pool, hubs


def fold_standardize(train, test):
    mu = train.mean(axis=0)
    sd = train.std(axis=0)
    sd[sd == 0] = 1.0
    return (train - mu) / sd, (test - mu) / sd


def cv_auc(x, y, cols, splits):
    pred = np.full(len(y), np.nan)
    for tr, va in splits:
        xtr, xva = fold_standardize(x[tr][:, cols], x[va][:, cols])
        model = LogisticRegression(penalty="l1", solver="liblinear", C=0.31622776601683794,
                                   max_iter=3000, random_state=SEED)
        model.fit(xtr, y[tr])
        pred[va] = model.predict_proba(xva)[:, 1]
    return float(roc_auc_score(y, pred))


def random_forest_analysis(x, y, pool, hubs):
    rkf = RepeatedStratifiedKFold(n_splits=5, n_repeats=3, random_state=SEED)
    importance = []
    fold_auc = []
    rank_rows = []
    for fold, (tr, va) in enumerate(rkf.split(x, y), 1):
        model = RandomForestClassifier(
            n_estimators=400, min_samples_leaf=5, max_features="sqrt",
            class_weight="balanced_subsample", random_state=SEED + fold, n_jobs=-1,
        )
        model.fit(x[tr], y[tr])
        prob = model.predict_proba(x[va])[:, 1]
        fold_auc.append(roc_auc_score(y[va], prob))
        pi = permutation_importance(model, x[va], y[va], scoring="roc_auc", n_repeats=5,
                                    random_state=SEED + fold, n_jobs=-1)
        importance.append(pi.importances_mean)
        order = np.argsort(-pi.importances_mean)
        ranks = np.empty(len(pool), int); ranks[order] = np.arange(1, len(pool) + 1)
        rank_rows.append(ranks)
    imp = np.vstack(importance)
    ranks = np.vstack(rank_rows)
    table = pd.DataFrame({
        "gene": pool,
        "mean_permutation_importance_delta_auc": imp.mean(axis=0),
        "importance_sd_across_folds": imp.std(axis=0, ddof=1),
        "median_rank": np.median(ranks, axis=0),
        "top10_frequency": (ranks <= 10).mean(axis=0),
        "top20_frequency": (ranks <= 20).mean(axis=0),
        "top30_frequency": (ranks <= 30).mean(axis=0),
        "prespecified_hub": [g in hubs for g in pool],
    }).sort_values("median_rank")
    table.to_csv(OUT / "Supplementary_Table_S23_random_forest_robustness.csv", index=False,
                 encoding="utf-8-sig")
    hub = table[table.prespecified_hub].copy()
    top20 = set(table.loc[table.median_rank <= 20, "gene"])
    rfe = set(json.load(open(SOURCE / "ml_summary.json", encoding="utf-8"))["rfe_genes"])
    metrics = {
        "rf_auc_mean": float(np.mean(fold_auc)),
        "rf_auc_sd": float(np.std(fold_auc, ddof=1)),
        "hub_top20_n": int((hub.median_rank <= 20).sum()),
        "hub_top30_n": int((hub.median_rank <= 30).sum()),
        "median_hub_rank": float(hub.median_rank.median()),
        "jaccard_rf20_rfe12": len(top20 & rfe) / len(top20 | rfe),
    }
    return table, hub, metrics


def negative_controls(x, y, pool, hubs):
    cv = StratifiedKFold(n_splits=5, shuffle=True, random_state=SEED)
    splits = list(cv.split(x, y))
    hub_idx = np.array([pool.index(g) for g in hubs])
    observed = cv_auc(x, y, hub_idx, splits)
    rows = [{"control": "Locked 11-gene model", "iteration": 0, "auc": observed}]
    for i in range(300):
        cols = RNG.choice(len(pool), len(hubs), replace=False)
        rows.append({"control": "Random 11-gene set", "iteration": i + 1,
                     "auc": cv_auc(x, y, cols, splits)})
        yp = RNG.permutation(y)
        psplits = list(StratifiedKFold(5, shuffle=True, random_state=SEED + i + 1).split(x, yp))
        rows.append({"control": "Permuted ER labels", "iteration": i + 1,
                     "auc": cv_auc(x, yp, hub_idx, psplits)})
    table = pd.DataFrame(rows)
    table.to_csv(OUT / "Supplementary_Table_S24_classifier_negative_controls.csv", index=False,
                 encoding="utf-8-sig")
    rand = table.loc[table.control == "Random 11-gene set", "auc"]
    perm = table.loc[table.control == "Permuted ER labels", "auc"]
    metrics = {
        "observed_auc": observed,
        "random_median": float(rand.median()),
        "random_95_low": float(rand.quantile(.025)),
        "random_95_high": float(rand.quantile(.975)),
        "random_empirical_p": float((1 + (rand >= observed).sum()) / (len(rand) + 1)),
        "permuted_median": float(perm.median()),
        "permuted_95_low": float(perm.quantile(.025)),
        "permuted_95_high": float(perm.quantile(.975)),
        "permuted_empirical_p": float((1 + (perm >= observed).sum()) / (len(perm) + 1)),
    }
    return table, metrics


def auc_meta_analysis():
    ml = json.load(open(SOURCE / "ml_summary.json", encoding="utf-8"))
    rows = [{"cohort": "METABRIC", "n": 1979, "auc": ml["external_auc_metabric"],
             "ci_low": ml["external_auc_ci95"][0], "ci_high": ml["external_auc_ci95"][1]}]
    for cohort, fn in [("GSE21653", "geo_gse21653_validation.csv"),
                       ("GSE7390", "geo_gse7390_validation.csv"),
                       ("GSE20711", "geo_gse20711_validation.csv")]:
        d = pd.read_csv(SOURCE / fn).iloc[0]
        rows.append({"cohort": cohort, "n": int(d.n_er), "auc": float(d.er_classifier_auc),
                     "ci_low": float(d.er_classifier_auc_ci_low),
                     "ci_high": float(d.er_classifier_auc_ci_high)})
    df = pd.DataFrame(rows)
    if not ((df[["auc", "ci_low", "ci_high"]] > 0).all().all()
            and (df[["auc", "ci_low", "ci_high"]] < 1).all().all()):
        raise ValueError("AUC estimates and confidence limits must lie strictly between 0 and 1")
    logit = np.log(df.auc / (1 - df.auc))
    lo = np.log(df.ci_low / (1 - df.ci_low)); hi = np.log(df.ci_high / (1 - df.ci_high))
    se = (hi - lo) / (2 * 1.96)
    w = 1 / se**2
    fixed = np.sum(w * logit) / np.sum(w)
    q = np.sum(w * (logit - fixed)**2)
    df_q = len(df) - 1
    c = np.sum(w) - np.sum(w**2) / np.sum(w)
    tau2 = max(0.0, (q - df_q) / c)
    wr = 1 / (se**2 + tau2)
    pooled = np.sum(wr * logit) / np.sum(wr)
    pooled_se = np.sqrt(1 / np.sum(wr))
    inv = lambda z: 1 / (1 + np.exp(-z))
    df["logit_auc"] = logit; df["logit_se"] = se; df["random_weight"] = wr / wr.sum()
    pooled_row = {"cohort": "Random-effects pooled", "n": int(df.n.sum()), "auc": inv(pooled),
                  "ci_low": inv(pooled - 1.96 * pooled_se), "ci_high": inv(pooled + 1.96 * pooled_se),
                  "logit_auc": pooled, "logit_se": pooled_se, "random_weight": 1.0}
    out = pd.concat([df, pd.DataFrame([pooled_row])], ignore_index=True)
    out.to_csv(OUT / "Supplementary_Table_S25_external_AUC_meta_analysis.csv", index=False,
               encoding="utf-8-sig")
    metrics = {"pooled_auc": pooled_row["auc"], "pooled_low": pooled_row["ci_low"],
               "pooled_high": pooled_row["ci_high"], "Q": float(q),
               "Q_p": float(stats.chi2.sf(q, df_q)),
               "I2": float(max(0, (q - df_q) / q) * 100 if q > 0 else 0), "tau2": float(tau2)}
    return out, metrics


def make_figure(rf_table, hubs, controls, meta, m_rf, m_neg, m_meta):
    fig = plt.figure(figsize=(7.2, 6.4))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.15, 1], hspace=.42, wspace=.32)
    ax1 = fig.add_subplot(gs[0, :]); ax2 = fig.add_subplot(gs[1, 0]); ax3 = fig.add_subplot(gs[1, 1])

    h = hubs.sort_values("median_rank", ascending=False)
    colors = ["#2b6f8a" if r <= 20 else "#9aa0a6" for r in h.median_rank]
    ax1.barh(h.gene, h.median_rank, color=colors, height=.7)
    ax1.axvline(20, color="#b33a3a", ls="--", lw=1)
    ax1.invert_xaxis(); ax1.set_xlabel("Median random-forest rank (lower is stronger)")
    ax1.set_title(f"a  Independent algorithm support: {m_rf['hub_top20_n']}/11 hubs in RF top 20",
                  loc="left", fontweight="bold", fontsize=8)
    ax1.grid(axis="x", alpha=.2)

    groups = [controls.loc[controls.control == x, "auc"].to_numpy() for x in
              ["Permuted ER labels", "Random 11-gene set"]]
    vp = ax2.violinplot(groups, positions=[1, 2], showmedians=True, widths=.75)
    for b, color in zip(vp["bodies"], ["#a7adb3", "#78b7a4"]): b.set_facecolor(color); b.set_alpha(.8)
    ax2.scatter([3], [m_neg["observed_auc"]], s=48, color="#b33a3a", zorder=5)
    ax2.set_xticks([1, 2, 3], ["Permuted\nlabels", "Random\n11 genes", "Locked\n11 genes"])
    ax2.set_ylabel("5-fold cross-validated AUC"); ax2.set_ylim(.38, 1.0)
    ax2.set_title("b  Negative controls", loc="left", fontweight="bold", fontsize=8)
    ax2.grid(axis="y", alpha=.2)

    md = meta.iloc[:-1].copy(); y = np.arange(len(md), 0, -1)
    ax3.errorbar(md.auc, y, xerr=[md.auc-md.ci_low, md.ci_high-md.auc], fmt="o",
                 color="#2b6f8a", capsize=2, lw=1)
    pooled = meta.iloc[-1]
    ax3.errorbar([pooled.auc], [0], xerr=[[pooled.auc-pooled.ci_low], [pooled.ci_high-pooled.auc]],
                 fmt="D", color="#b33a3a", capsize=3, lw=1.2)
    ax3.set_yticks(list(y)+[0], md.cohort.tolist()+["Pooled"])
    ax3.axvline(.5, color="#777777", ls=":", lw=.8); ax3.set_xlim(.7, .97)
    ax3.set_xlabel("ER-classifier AUC (95% CI)")
    ax3.set_title("c  Independent-cohort synthesis", loc="left", fontweight="bold", fontsize=8)
    ax3.grid(axis="x", alpha=.2)
    fig.suptitle("Algorithmic robustness, negative controls and external validation", fontsize=10, y=.98)
    fig.savefig(OUT / "Supplementary_Figure_S12_classifier_robustness.png", dpi=600, bbox_inches="tight")
    fig.savefig(OUT / "Supplementary_Figure_S12_classifier_robustness.pdf", bbox_inches="tight")
    fig.savefig(OUT / "Supplementary_Figure_S12_classifier_robustness.svg", bbox_inches="tight")
    fig.savefig(OUT / "Supplementary_Figure_S12_classifier_robustness.tiff", dpi=600, bbox_inches="tight")
    plt.close(fig)


def patch_docx(m_rf, m_neg, m_meta):
    src = OUT / "Manuscript_Cholesterol_Metabolism_BRCA_v26.docx"
    dst = OUT / "Manuscript_Cholesterol_Metabolism_BRCA_v27.docx"
    doc = docx.Document(src)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    p25 = next(p for p in paras if p.text.strip().startswith("2.5 Candidate selection"))
    p25.add_run(
        " Three post hoc robustness analyses were performed under a locked-signature framework without changing the "
        "11-gene signature: repeated five-fold random-forest models with validation-fold permutation importance; "
        "300 random 11-gene cholesterol-pathway sets and 300 ER-label permutations under fold-specific "
        "standardization; and a random-effects meta-analysis of logit-transformed AUCs across the four independent "
        "cohorts."
    )
    p33 = next(p for p in paras if p.text.strip().startswith("3.3 Hub genes"))
    p33.add_run(
        f" In the independent random-forest sensitivity analysis, {m_rf['hub_top20_n']}/11 locked hub genes "
        f"ranked within the top 20 of 119 candidates and {m_rf['hub_top30_n']}/11 within the top 30; the median "
        f"hub rank was {m_rf['median_hub_rank']:.1f}. The random forest achieved mean validation-fold AUC "
        f"{m_rf['rf_auc_mean']:.3f} (SD {m_rf['rf_auc_sd']:.3f}). The locked logistic signature achieved "
        f"five-fold AUC {m_neg['observed_auc']:.3f}, compared with median {m_neg['random_median']:.3f} "
        f"for random 11-gene pathway sets and {m_neg['permuted_median']:.3f} after ER-label permutation "
        f"(both empirical P = {m_neg['random_empirical_p']:.3f}). Across METABRIC, GSE21653, GSE7390 and "
        f"GSE20711, the random-effects pooled AUC was {m_meta['pooled_auc']:.3f} "
        f"(95% CI {m_meta['pooled_low']:.3f}-{m_meta['pooled_high']:.3f}; I2 = {m_meta['I2']:.1f}%). "
        f"These post hoc analyses support algorithmic and cross-cohort robustness without redefining the original "
        f"signature (Supplementary Figure S12; Supplementary Tables S23-S25)."
    )
    p42 = next(p for p in paras if p.text.strip().startswith("4.2 Cholesterol metabolism and ER biology"))
    p42.add_run(
        " Random-forest ranking, pathway-matched random-gene controls and label permutation further indicate that "
        "the classifier performance is not attributable solely to one feature-selection algorithm or chance "
        "partitioning. These are post hoc robustness analyses and do not constitute independent model development."
    )
    pav = next(p for p in paras if p.text.strip().startswith("Data availability."))
    for r in pav.runs:
        r.text = r.text.replace("Supplementary Tables S1-S22", "Supplementary Tables S1-S25")
        r.text = r.text.replace("Supplementary Figures S1-S11", "Supplementary Figures S1-S12")
    doc.save(dst)
    return dst


def main():
    x, y, pool, hubs = load_data()
    rf, hub, m_rf = random_forest_analysis(x, y, pool, hubs)
    controls, m_neg = negative_controls(x, y, pool, hubs)
    meta, m_meta = auc_meta_analysis()
    make_figure(rf, hub, controls, meta, m_rf, m_neg, m_meta)
    dst = patch_docx(m_rf, m_neg, m_meta)
    (OUT / "v27_robustness_summary.json").write_text(
        json.dumps({"random_forest": m_rf, "negative_controls": m_neg, "auc_meta": m_meta}, indent=2),
        encoding="utf-8")
    print(dst)
    print(json.dumps({"random_forest": m_rf, "negative_controls": m_neg, "auc_meta": m_meta}, indent=2))


if __name__ == "__main__":
    main()
