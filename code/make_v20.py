# -*- coding: utf-8 -*-
"""Generate submission-ready v20 from v15.

Changes (on top of v19):
0. Fourth independent validation cohort GSE20711 (Fackler et al. 2011,
   n = 88 with RFS; ER-classifier AUC 0.833) added; new reference [43],
   old 43-52 renumbered to 44-53.

Changes (v19, on top of v18):
1. LASSO bootstrap stability (300 resamples) added to Section 3.3.
2. CIBERSORT/ESTIMATE deconvolution + immune-checkpoint genes added to 3.5.
3. GEO subtype external validation (GSE21653/GSE7390) added after the
   consensus-clustering paragraph (Supplementary Figure S3).
4. HPA protein evidence + pan-cancer FDPS expression added to Section 3.10.
5. Published lipid-signature comparison table (Supplementary Table S11) and
   TRIPOD-AI checklist (S12) referenced; five new references appended.
6. Limitation 4.5 updated (CIBERSORT/ESTIMATE now performed; GSE7390 subtype
   survival non-significant noted; references 40/41 re-cited in Section 3.5).
7. Data/code availability updated to S11-S12 and graphical abstract.
"""
import re
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = Path(r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature")
V15 = SRC / "Manuscript_Cholesterol_Metabolism_BRCA_v15.docx"
V16 = SRC / "Manuscript_Cholesterol_Metabolism_BRCA_v16.docx"
V17 = SRC / "Manuscript_Cholesterol_Metabolism_BRCA_v17.docx"
V18 = SRC / "Manuscript_Cholesterol_Metabolism_BRCA_v18.docx"
V19 = SRC / "Manuscript_Cholesterol_Metabolism_BRCA_v19.docx"
V20 = SRC / "Manuscript_Cholesterol_Metabolism_BRCA_v20.docx"

GITHUB_URL = "https://github.com/y0307Yd/brca-cholesterol-signature"
ZENODO_DOI = "https://doi.org/10.5281/zenodo.21873168"


ABSTRACT = [
    ("Purpose. ",
     "Cholesterol metabolism is closely linked to estrogen-receptor (ER) biology, "
     "but its value for breast cancer classification, molecular subtyping and "
     "causal inference remains incompletely quantified."),
    ("Materials and Methods. ",
     "We applied co-expression network screening, machine-learning feature "
     "selection, consensus clustering, single-cell localization, immune "
     "deconvolution (CIBERSORT/ESTIMATE), SMR/HEIDI Mendelian randomization "
     "and Bayesian colocalisation to TCGA-BRCA (n = 952) and METABRIC "
     "(n = 1,979), with classifier validation in GSE21653, GSE7390 and "
     "GSE20711 and nearest-centroid subtype mapping across all external "
     "cohorts."),
    ("Results. ",
     "The strict WGCNA-by-pathway intersection was empty. A pathway-first "
     "adaptation using 119 measurable cholesterol genes yielded 11 hub genes; "
     "the ER classifier reached AUCs of 0.946 (5-fold cross-validation), 0.922 "
     "(METABRIC), 0.847 (GSE21653) and 0.903 (GSE7390). Consensus clustering "
     "defined four subtypes with divergent ER composition, immune infiltration "
     "and PAM50 correspondence. Single-cell analyses localized cholesterol "
     "biosynthesis to tumour epithelial and myeloid/endothelial compartments. "
     "Pathway-wide SMR identified FDPS as a putatively causal protective gene "
     "(P = 9.8e-7; HEIDI P = 0.026; coloc.abf PP.H4 = 0.9985), with the "
     "protective direction replicated in ER+ OncoArray analyses (P = 1.54e-5) "
     "and an inverse methylation-expression correlation (ER-adjusted partial "
     "rho = -0.187, P = 2.2e-6)."),
    ("Conclusion. ",
     "Cholesterol metabolism genes do not define a progression-associated "
     "co-expression module but carry a strong, cross-platform ER signature, "
     "define immune-relevant molecular subtypes and include FDPS as a "
     "putatively causal protective gene for breast cancer risk."),
]

KEYWORDS = ("Keywords: Breast Neoplasms; Cholesterol; Machine Learning; "
            "Receptors, Estrogen; Tumor Microenvironment")

# key: unique prefix of the v15 paragraph -> (bold part, body part)
# A trailing "\\n" in the body part inserts a line break.
BODY = {
    "Cholesterol and its derivatives support":
        ("",
         "Cholesterol and its derivatives support breast cancer cell "
         "proliferation, membrane biogenesis, steroid hormone synthesis and "
         "immune regulation [1,2]. ER-positive (ER+) tumours coordinately "
         "up-regulate cholesterol biosynthetic enzymes, and cholesterol-lowering "
         "agents have been proposed as adjunct therapy [3-6]. Most "
         "lipid-metabolism transcriptomic studies have focused on prognosis "
         "[7,8]; here we additionally report an honest null network-pathway "
         "intersection, an ER-focused classifier validated on three platforms, "
         "and a germline causal-inference layer."),
    "A widely used bioinformatics template":
        ("",
         "We applied a five-layer bioinformatics framework - co-expression "
         "network screening, machine-learning feature selection, patient-level "
         "molecular subtyping, single-cell localization and Mendelian-"
         "randomization causal testing - to cholesterol metabolism in breast "
         "cancer, using TCGA-BRCA and METABRIC as discovery and external "
         "cohorts."),
    "2.1 Data sources and verification":
        ("2.1 Data sources and verification ",
         "The overall study design and the flow of analyses are summarised in "
         "Supplementary Figure S2. For validation analyses, PAM50 calls for "
         "TCGA-BRCA were obtained from "
         "UCSC Xena (GDC hub) [9]; FDPS HM450 methylation from the cBioPortal "
         "data repository (brca_tcga_methylation_hm450); the independent "
         "expression cohorts GSE21653 (Affymetrix HG-U133 Plus 2.0, NCBI GEO "
         "with GPL570 annotation) [10] and GSE7390 (Affymetrix HG-U133A, "
         "ArrayExpress E-GEOD-7390 processed matrix and SDRF clinical data, "
         "probes mapped with hgu133a.db) [11] were used; and the second "
         "single-cell atlas GSE161529 (Mendeley Data mirror of the Pal et al. "
         "atlas) [12] was used for single-cell validation. A fourth "
         "Affymetrix cohort (GSE20711, NCBI GEO) was used for classifier "
         "transfer. TCGA-BRCA "
         "RNA-sequencing STAR counts, methylation, mutation and clinical files "
         "were obtained from the NCI GDC portal [13] and verified against the "
         "official manifest by MD5 checksums (4,393/4,393 files); progression-"
         "free interval (PFI) was taken from TCGA-CDR [14]. METABRIC expression "
         "and clinical annotation with recurrence-free survival (RFS) were "
         "obtained from cBioPortal [15]."),
    "The discovery cohort comprised":
        ("",
         "The discovery cohort comprised 952 TCGA-BRCA patients [16] with valid "
         "PFI (122 events, 12.8%; median follow-up 2.1 years; ER status for "
         "909). The external cohort comprised 1,979 METABRIC samples [17] with "
         "RFS (803 events, 40.6%; median follow-up 8.5 years; ER status for "
         "all)."),
    "A cholesterol metabolism gene set":
        ("",
         "A cholesterol metabolism gene set was assembled from KEGG hsa00100 "
         "(steroid biosynthesis) [18] and Gene Ontology biological-process terms "
         "GO:0008203 (cholesterol metabolic process) and GO:0006695 "
         "(cholesterol biosynthetic process) [19] using org.Hs.eg.db [20]; the "
         "union contained 120 unique genes (19 from KEGG, 107 from GO:0008203, "
         "39 from GO:0006695)."),
    "Expression was transformed as":
        ("",
         "Expression was transformed as log2(CPM+1). The 8,000 most variable "
         "genes entered weighted correlation network analysis [21,22]: "
         "soft-threshold power selection by scale-free fit, topological "
         "overlap matrix, average-linkage hierarchical clustering and merging "
         "of modules with eigengene correlation >= 0.80. Module eigengenes "
         "were correlated with PFI, ER status and stage with "
         "Benjamini-Hochberg FDR [23]."),
    "2.5 Candidate selection and machine learning":
        ("2.5 Candidate selection and machine learning ",
         "The strict intersection of trait-associated modules with the "
         "cholesterol set was empty; the candidate pool was therefore the 119 "
         "cholesterol genes measurable in both cohorts (pathway-first "
         "adaptation, documented as an analysis deviation). LASSO logistic "
         "regression [24] and SVM-RFE (linear SVM) [25] were applied to ER "
         "status; the intersection of their selected lists defined 11 hub "
         "genes. Selection stability was assessed with 300 bootstrap "
         "resamples of the LASSO fit. A signature score was computed as the "
         "sum of LASSO "
         "coefficients times z-scored expression. Performance was assessed by "
         "5-fold cross-validation (CV) AUC in TCGA and external AUC in "
         "METABRIC (2,000-bootstrap 95% confidence intervals); calibration was "
         "evaluated with reliability curves, Brier scores and logistic "
         "recalibration; clinical utility with decision-curve analysis."),
    "Consensus clustering [26]":
        ("",
         "Consensus clustering [26] (100 resamples, 80% of patients, k-means, "
         "k = 2-6) was applied to hub-gene z-scores; the number of clusters was "
         "selected by the proportion of ambiguous clustering (PAC). Subtypes "
         "were compared for ER composition, survival, immune marker scores "
         "(nine curated cell-type panels), the cholesterol-pathway score, "
         "differential expression (Mann-Whitney, BH-FDR) and KEGG/GO "
         "enrichment of strong-effect genes. TCGA subtype centroids were "
         "mapped onto METABRIC, GSE21653 and GSE7390 by nearest-centroid "
         "assignment. Immune programs were quantified by single-sample GSEA "
         "(ssGSEA) [27], CIBERSORT (nu-SVR, LM22, 22 immune cell types) [28] "
         "and ESTIMATE (immune/stromal scores) [29]; immune-checkpoint genes "
         "(CD274, PDCD1, CTLA4, LAG3, HAVCR2, IDO1, CD8A) were compared "
         "across subtypes. Hub-gene and FDPS expression were compared across "
         "PAM50 intrinsic subtypes."),
    "Python 3.11":
        ("",
         "Python 3.11 (scikit-learn [30], scipy, lifelines [31], pandas, numpy, "
         "matplotlib) and R 4.6.1 (org.Hs.eg.db, hgu133a.db, susieR, coloc). "
         "All random seeds fixed at 20260804."),
    "2.8 Single-cell localization and SMR causal inference":
        ("2.8 Single-cell localization and SMR causal inference\n",
         "Single-cell localization used the processed GSE176078 atlas [32] "
         "(100,064 cells; CELLxGENE release; three annotation resolutions); "
         "hub-gene mean expression, detection percentage and Wilcoxon "
         "enrichment (BH-FDR) were computed per cell type and confirmed in "
         "GSE161529 [12]. Causal inference used SMR/HEIDI v1.3.1 [33,34] with "
         "eQTLGen whole-blood cis-eQTLs [35] (31,684 individuals), the BCAC "
         "overall and triple-negative breast cancer GWAS [36] (247,173 and "
         "118,987 women) and the 1000 Genomes Phase 3 European LD reference, "
         "scanning 119 cholesterol genes (251 probes) with a +/-2 Mb cis "
         "window, instrument eQTL P < 5e-8 and Bonferroni alpha = 0.05/119. "
         "ER-stratified sensitivity used the BCAC OncoArray public release "
         "[37] (ER+ 69,501; ER- 21,468 cases). GTEx v8 breast-mammary and "
         "whole-blood SMR-format BESDs [38,39] were used for tissue-specific "
         "and replication analyses. For FDPS, single-causal-variant Bayesian "
         "colocalisation (coloc.abf; p1 = p2 = 1e-4, p12 = 1e-5, W = 0.2) and "
         "multi-signal SuSiE/coloc.susie [40,41] were applied; detailed "
         "parameter settings are provided in the archived code."),
    "TCGA-BRCA contributed 952":
        ("",
         "TCGA-BRCA contributed 952 patients (122 PFI events; median follow-up "
         "2.1 years; ER+ 699/909). METABRIC contributed 1,979 samples (803 RFS "
         "events; median follow-up 8.5 years; ER+ 1,505/1,979). Both cohorts "
         "covered 19,448 shared genes."),
    "At soft-threshold power":
        ("",
         "At soft-threshold power beta = 5 (scale-free fit R2 = 0.834), the "
         "8,000-gene network resolved into 15 modules (38-853 genes). The "
         "module most strongly associated with PFI events (M2238; 69 genes) "
         "showed a weak correlation (r = 0.082, P = 0.012) that did not "
         "survive multiple-testing correction (BH-FDR = 0.170), whereas ER "
         "status showed very strong module associations (M2314; r = -0.794, "
         "BH-FDR ~ 4e-197)."),
    "Strictly intersecting any trait-associated module":
        ("",
         "The strict intersection of trait-associated modules with the "
         "120-gene cholesterol set was empty: the PFI module (M2238) and the "
         "ER module (M2314) contained no cholesterol genes. No cholesterol "
         "gene was significantly associated with PFI (smallest Cox P = 0.13, "
         "CH25H), whereas several strongly tracked ER status (LIPE "
         "P = 4.5e-17; EBP P = 2.4e-12; FDXR P = 6.0e-12). This honest null "
         "result motivated the pathway-first adaptation below."),
    "3.3 Hub genes and the ER-status classifier":
        ("3.3 Hub genes and the ER-status classifier ",
         "From 119 measurable cholesterol genes, LASSO (C = 0.316; 69 genes) "
         "and SVM-RFE (12 genes) selected 11 hub genes by intersection "
         "(Table 3). In 300 LASSO bootstrap resamples all 11 hub genes were "
         "selected in at least 73.7% of fits (mean 90.2%; G6PD and VLDLR "
         "100%, PRKAA1 99.7%, LIMA1 73.7%), confirming stability. "
         "The ER-status classifier reached internal 5-fold CV AUC "
         "0.946 +/- 0.012 (full-data 0.953) and external AUC 0.922 "
         "(95% CI 0.907-0.935) in METABRIC. Calibration was excellent in TCGA "
         "out-of-fold predictions (Brier 0.069; recalibration intercept 0.002, "
         "slope 1.06) and mildly underconfident in METABRIC (Brier 0.098; "
         "intercept 0.57, slope 1.41); decision-curve analysis showed positive "
         "net benefit across thresholds (Figure 3). The signature was not "
         "associated with PFI in TCGA (adjusted HR 0.90, P = 0.151) but showed "
         "an ER/stage-adjusted protective association with RFS in METABRIC "
         "(HR 0.90, 95% CI 0.82-0.98, P = 0.013)."),
    "Consensus clustering of hub-gene expression":
        ("",
         "Consensus clustering selected k = 4 (PAC = 0.097). Subtypes differed "
         "markedly in ER composition (C1 13.6%, C2 93.8%, C3 97.5%, C4 60.1% "
         "ER+), signature score and cholesterol-pathway score (Table 4); TCGA "
         "survival differences were not significant (multivariate log-rank "
         "P = 0.331). Nearest-centroid mapping onto METABRIC reproduced the "
         "four subtypes with near-identical ER composition and significant "
         "survival differences (global log-rank P = 7.8e-05; C4 worst, C2 "
         "best), which attenuated after adjustment for ER status and stage "
         "(C4 vs C1 HR = 1.26, 95% CI 0.95-1.67, P = 0.12), indicating that "
         "the unadjusted differences largely tracked ER status and stage "
         "(Figure 7). Nearest-centroid mapping onto GSE21653 (n = 252) and "
         "GSE7390 (n = 198) reproduced the subtype ER gradient "
         "(Kruskal-Wallis P = 4.2e-16 and P = 4.7e-16, respectively); "
         "survival differences were significant in GSE21653 (global log-rank "
         "P = 0.026; C1 worst, C2 best) but not in GSE7390 (P = 0.48), "
         "consistent with ER-driven survival differences (Supplementary "
         "Figure S3)."),
    "3.5 Immune microenvironment and pathway programs":
        ("3.5 Immune microenvironment and pathway programs across subtypes\n",
         "All nine immune/stromal marker scores differed across subtypes "
         "(Kruskal-Wallis P < 0.001): C1 showed the highest CD8 T, regulatory "
         "T, NK, B-cell and macrophage scores and the lowest stromal score; "
         "C3 the lowest across most immune populations; C2 the highest "
         "dendritic-cell and stromal scores; C4 the highest neutrophil score; "
         "the pattern reproduced in METABRIC (Supplementary Figure S1). "
         "Strong-effect differential genes (|mean log2 difference| >= 1: "
         "1,840/304/225/92 for C1-C4) showed C1-enriched cell-cycle programs "
         "and C3-enriched oxidative-phosphorylation and ribosomal programs. "
         "ssGSEA confirmed these patterns (all Kruskal-Wallis P < 0.05 except "
         "B cells, P = 0.42) with marker-score correlations of 0.40-0.93 "
         "(Figure 8; Supplementary Table S8). ESTIMATE confirmed the "
         "immune-hot/immune-cold axis (ImmuneScore highest in C1, "
         "P = 1.9e-3; StromalScore highest in C2, P = 4.2e-22; "
         "ESTIMATEScore P = 3.5e-6). CIBERSORT showed higher C1 fractions of "
         "M1/M0 macrophages, activated dendritic cells, follicular helper T "
         "cells, monocytes and activated NK cells, and higher C2/C3 "
         "fractions of resting mast cells and M2 macrophages (all P < "
         "0.001). Immune-checkpoint genes CD274, PDCD1, CTLA4, LAG3 and IDO1 "
         "were highest in C1 (P = 1.7e-20 to 4.9e-5); HAVCR2 and CD8A did "
         "not differ (P = 0.12 and 0.11; Supplementary Tables S8 and S14)."),
    "3.6 Single-cell localization of hub genes":
        ("3.6 Single-cell localization of hub genes\n",
         "In the 100,064-cell atlas, cholesterol biosynthetic enzymes "
         "(DHCR24, DHCR7, G6PD, HSD17B7, FDXR) were most strongly enriched in "
         "malignant epithelial cells (all Wilcoxon BH-FDR < 1e-10), while "
         "ABCG1, G6PD and LIMA1 were enriched in myeloid/endothelial cells "
         "and cancer-associated fibroblasts, localizing the pathway to both "
         "tumour and microenvironment compartments."),
    "3.7 SMR causal inference":
        ("3.7 SMR causal inference\n",
         "Pathway-wide SMR identified three probes passing the Bonferroni "
         "threshold in overall breast cancer: CYP51A1 (P = 1.4e-8), FDPS "
         "(P = 9.8e-7) and LSS (P = 4.1e-5). HEIDI rejected CYP51A1 and LSS as "
         "LD-confounded but supported FDPS (P_HEIDI = 0.026), with higher "
         "genetically predicted FDPS expression associated with lower breast "
         "cancer risk (b = -0.42, 95% CI -0.59 to -0.25); SREBF1, LIPA and "
         "FAXDC2 showed supportive protective signals. In the OncoArray "
         "ER-stratified sensitivity analysis the protective direction "
         "replicated (ER+ P = 1.54e-05, b = -0.44; ER- P = 0.28, b = -0.15, "
         "underpowered; lead SNP rs6677385), and the OncoArray overall "
         "analysis gave P = 9.02e-06. TNBC sensitivity found no "
         "Bonferroni-significant probe; G6PD, HMGCS2 and NSDHL lacked "
         "eQTLGen instruments. Per-probe results are provided in "
         "Supplementary Tables S2 and S10."),
    "All nine immune/stromal marker scores differed":
        ("", ""),
    "From 119 measurable cholesterol genes, LASSO (C = 0.316) selected":
        ("", ""),
    "For single-cell localization (Layer 4)":
        ("", ""),
    "TCGA-BRCA RNA-sequencing STAR counts":
        ("", ""),
    "Multi-signal colocalisation (SuSiE":
        ("", ""),
    "In an independent Affymetrix GPL570 cohort":
        ("", ""),
    "In the independent GSE161529 atlas":
        ("", ""),
    "In 672 TCGA-BRCA tumours with HM450":
        ("", ""),
    "3.8 GTEx breast-tissue SMR sensitivity":
        ("3.8 GTEx breast-tissue SMR sensitivity analysis\n",
         "With GTEx v8 breast-mammary cis-eQTLs, 109/119 cholesterol genes had "
         "probes and 16 had genome-wide-significant instruments; no probe "
         "passed Bonferroni in overall or TNBC analyses. LIPA (P = 8.2e-3; "
         "P_HEIDI = 0.53) and FAXDC2 (P = 1.15e-2; P_HEIDI = 0.69) replicated "
         "direction-consistently with HEIDI support, matching the eQTLGen "
         "results. FDPS lacked a genome-wide-significant breast-tissue "
         "instrument, indicating that the blood-based FDPS signal is likely "
         "mediated through non-mammary tissues."),
    "3.9 Colocalisation supports":
        ("3.9 Colocalisation supports a shared causal variant at the FDPS locus\n",
         "coloc.abf over 3,248 shared SNPs gave PP.H4 = 0.9985 at the FDPS "
         "locus (window sensitivity 0.996-0.997 at +/-1 Mb, +/-500 kb and "
         "+/-250 kb). The lead SNP rs12091730 was strongly associated with "
         "both FDPS expression (P = 2.5e-18) and breast cancer risk "
         "(P = 1.8e-12) with opposing directions, forming one LD block (11 "
         "SNPs, r2 >= 0.8). Single-SNP conditional analysis confirmed a single "
         "eQTL signal and additional GWAS-only signals (rs4971059, "
         "rs11264454) with no eQTL effects; sensitivity at p12 = 1e-8 gave "
         "PP.H4 = 0.40. However, multi-signal SuSiE/coloc.susie identified "
         "four eQTL credible sets and one GWAS credible set with maximum "
         "PP.H4 = 0.44 (PP.H3 = 0.56), so the single-shared-variant "
         "interpretation of coloc.abf should be regarded as suggestive "
         "(Figure 9; Supplementary Table S1). Bayesian colocalisation was "
         "extended to LIPA, FAXDC2 and SREBF1; none supported a shared "
         "causal variant with breast cancer risk (PP.H4 = 0.035, 0.043 and "
         "0.171; PP.H1 dominated), indicating strong eQTL but no "
         "corresponding GWAS signal (Supplementary Table S15)."),
    "3.10 FDPS expression and clinical features":
        ("3.10 FDPS expression and clinical features\n",
         "FDPS expression was higher in ER-negative tumours (TCGA-BRCA "
         "P = 5.7e-17, Cohen's d = -0.69; METABRIC P = 2.2e-48, d = -0.79), "
         "unrelated to stage, and not independently prognostic after "
         "ER/stage adjustment (METABRIC adjusted HR = 1.03, P = 0.48). The "
         "germline eQTL-based causal signal is therefore distinct from the "
         "tumour-level association. In the Human Protein Atlas [42], FDPS showed "
         "cytoplasmic protein localization, tissue-enhanced expression in "
         "normal breast, and a non-significant TCGA protein-level prognosis "
         "(P = 0.195; HPA validation P = 0.023, classified unprognostic), "
         "consistent with the transcript-level null. FDPS mRNA was detected "
         "in all 17 queried cancer types (Supplementary Table S13)."),
    "3.11 GTEx v8 whole-blood sensitivity analysis":
        ("3.11 GTEx v8 whole-blood sensitivity analysis\n",
         "No genome-wide-significant FDPS cis-eQTL was available in GTEx v8 "
         "whole blood (n = 670); under sample-size scaling from the eQTLGen "
         "instrument (z = -9.65, N = 31,684), the expected GTEx whole-blood "
         "z-statistic is -1.40 (29% power at P < 0.05; <0.01% at P < 5e-8), so "
         "this resource cannot independently instrument FDPS."),
    "PAM50 calls were available for 871":
        ("",
         "PAM50 calls were available for 871 TCGA-BRCA tumours. The hub-gene "
         "subtypes differed markedly in PAM50 composition (chi-square = 820.8, "
         "P = 5.62e-168; C1 86.4% Basal, C2 67.4% LumA, C3 48.6% LumA). "
         "Hub-gene and FDPS expression differed across all five PAM50 subtypes "
         "(Kruskal-Wallis P < 0.001 for all 12 genes), with FDPS highest in "
         "Basal and HER2-enriched tumours, consistent with its ER-negative "
         "enrichment (Figure 10; Supplementary Tables S4 and S9)."),
    "3.13 Independent validation in GSE21653":
        ("3.13 Independent validation in GSE21653\n",
         "The ER classifier transferred to GSE21653 (n = 252, 83 "
         "disease-free-survival events) with AUC = 0.847 (95% CI 0.798-0.895; "
         "Brier 0.210), and FDPS remained strongly ER-associated (Mann-Whitney "
         "P = 1.24e-10; d = -0.86); FDPS survival associations attenuated "
         "after adjustment (HR = 1.16, P = 0.250). In a third independent "
         "cohort (GSE7390, n = 198, 91 RFS events), the classifier transferred "
         "with AUC = 0.903 (95% CI 0.848-0.949; Brier 0.119) and FDPS remained "
         "strongly ER-associated (P = 4.78e-05; d = -0.68), confirming "
         "cross-platform reproducibility (Figures 12 and 13; Supplementary "
         "Table S5). In a fourth independent cohort (GSE20711, n = 88 with "
         "RFS and 39 events; Fackler et al. [43]), the ER classifier "
         "transferred with AUC = 0.833 (95% CI 0.739-0.913; Brier 0.226) and "
         "FDPS remained ER-associated (Mann-Whitney P = 2.4e-3; d = -0.64) "
         "with no FDPS survival association (HR = 0.97, P = 0.86), "
         "confirming cross-platform transferability "
         "(Supplementary Table S16)."),
    "3.14 Second single-cell atlas confirms":
        ("3.14 Second single-cell atlas confirms FDPS enrichment in "
         "monocyte/macrophage populations\n",
         "In GSE161529 (70,419 cells after QC), FDPS was most strongly "
         "expressed in monocyte/macrophage populations (Monocyte mean 0.65). "
         "T-myeloid ligand-receptor co-expression was detected for 7 of 20 "
         "curated pairs, strongest IL1B-IL1R2 (11.2% of T/myeloid cells), "
         "CXCL9/10/11-CXCR3 (2.0-2.3%) and CSF1-CSF1R (1.7%) (Figure 14; "
         "Supplementary Table S6)."),
    "3.15 FDPS DNA methylation is inversely associated":
        ("3.15 FDPS DNA methylation is inversely associated with FDPS "
         "expression in TCGA-BRCA\n",
         "In 672 TCGA-BRCA tumours with HM450 methylation and RNA-seq data, "
         "FDPS methylation correlated inversely with FDPS expression "
         "(Spearman rho = -0.203, 95% CI -0.274 to -0.133, P = 1.04e-07). The "
         "association persisted after ER adjustment (partial rho = -0.187, "
         "P = 2.2e-6) and in ER+ tumours (rho = -0.206, P = 4.6e-6), with a "
         "consistent but non-significant direction in ER- tumours "
         "(rho = -0.142, P = 0.087; n = 147) (Figures 15 and 16; "
         "Supplementary Table S7)."),
    "4.1 Principal findings":
        ("4.1 Principal findings ",
         "The five-layer framework was applied with real data. The strict "
         "WGCNA-by-pathway intersection was empty; the pathway-first "
         "adaptation produced an 11-gene signature that discriminated ER "
         "status with strong cross-platform performance (TCGA CV AUC 0.946; "
         "METABRIC 0.922; GSE21653 0.847; GSE7390 0.903; GSE20711 0.833). "
         "Consensus clustering "
         "defined four subtypes with divergent ER composition and immune "
         "infiltration, reproduced in METABRIC, GSE21653 and GSE7390. "
         "Single-cell analysis localized "
         "the signature to tumour epithelial and myeloid/endothelial "
         "compartments, and pathway-wide SMR identified FDPS as a putatively "
         "causal protective gene for breast cancer risk, supported by "
         "ER-stratified replication, HEIDI and coloc.abf (PP.H4 = 0.9985) but "
         "weaker under multi-signal SuSiE (max PP.H4 = 0.44). Tumour FDPS "
         "expression was strongly ER-associated and not independently "
         "prognostic."),
    "This study applied a five-layer bioinformatics framework":
        ("",
         "GTEx v8 breast-mammary sensitivity directionally replicated LIPA and "
         "FAXDC2 but could not instrument FDPS in breast tissue; GTEx v8 whole "
         "blood lacked a significant FDPS cis-eQTL, consistent with the power "
         "available at n = 670."),
    "The ER-status signal is biologically coherent":
        ("",
         "The ER-status signal is biologically coherent: ER+ tumours "
         "up-regulate cholesterol biosynthetic enzymes to support steroid "
         "production and membrane turnover [3-5], and six hub genes are direct "
         "enzymes or regulators of cholesterol/steroid synthesis. The external "
         "AUC of 0.922 across platforms confirms that the signal is not a "
         "TCGA artifact. The near-null PFI associations indicate that "
         "cholesterol gene expression adds little independent prognostic "
         "information. FDPS methylation was inversely correlated "
         "with expression independently of ER status, and FDPS expression was "
         "highest in Basal and HER2-enriched intrinsic subtypes (Figure 10)."),
    "The four cholesterol-metabolism subtypes recapitulate":
        ("",
         "The four subtypes recapitulate known breast cancer biology: C1 is "
         "predominantly ER-negative, proliferative (cell-cycle up) and "
         "immune-hot (highest CD8 T, regulatory T, NK, B-cell and macrophage "
         "scores); C3 is almost uniformly ER+, immune-cold and "
         "oxidative-phosphorylation-high; C2 shows down-regulated cell-cycle "
         "programs and high stromal/dendritic scores; C4 is a mixed, "
         "neutrophil-high group. Cholesterol-synthesis-high subtypes are thus "
         "luminal-like and immune-cold, whereas cholesterol-synthesis-low "
         "subtypes are more heterogeneous, proliferative and immune-active "
         "[2,8]. ssGSEA provided a rank-based, non-parametric complement "
         "confirming these patterns."),
    "All TCGA files were MD5-verified":
        ("",
         "All TCGA files were MD5-verified against the GDC manifest; external "
         "validation used independent cohorts and platforms; feature selection "
         "used two independent algorithms; subtype stability was assessed by "
         "consensus clustering with PAC; multiple-testing corrections were "
         "applied throughout; and the null results of the strict template were "
         "reported rather than engineered away."),
    "4.5 Limitations":
        ("4.5 Limitations\n",
         "First, the strict WGCNA-by-pathway intersection was empty, and the "
         "candidate pool relied on the documented pathway-first adaptation. "
         "Second, an ER classifier built from expression partly rediscovers "
         "ER-pathway biology; its value is descriptive rather than mechanistic. "
         "Third, TCGA had only 122 PFI events, limiting subtype survival "
         "power; the METABRIC association (P = 4.2e-05 univariable; P = 0.013 "
         "after ER/stage adjustment) is the stronger prognostic evidence, and "
         "subtype survival differences were not significant in GSE7390 "
         "(global log-rank P = 0.48). Fourth, single-cell localization is "
         "descriptive and observational, and the HPA protein evidence is "
         "limited to descriptive localization with a non-significant "
         "protein-level prognosis. Fifth, SMR used blood (eQTLGen) "
         "cis-eQTLs; breast-tissue instruments were unavailable for FDPS, and "
         "tissue-specific causal effects cannot be excluded. The CYP51A1 and "
         "LSS signals were HEIDI-rejected and should not be interpreted as "
         "causal; the FDPS HEIDI P-value (0.026) was not adjusted across "
         "probes, and GTEx v8 whole blood (n = 670) could not independently "
         "instrument FDPS. Coloc.abf PP.H4 = 0.9985 was not robust under "
         "SuSiE-based multi-signal colocalisation (maximum PP.H4 = 0.44; the "
         "SuSiE credible-set estimates did not fully converge with the "
         "external 1000 Genomes LD reference), so the shared-variant "
         "interpretation should be regarded as suggestive; conditional "
         "analysis confirmed additional GWAS-only signals without residual "
         "eQTL effects. The causal claim is confined to germline-regulated "
         "blood expression rather than tumour expression."),
    "4.6 Causal inference completed":
        ("4.6 Causal inference ",
         "By SMR with eQTLGen and BCAC summary statistics, causal inference "
         "was completed for 81 cholesterol-gene probes; FDPS passed Bonferroni "
         "correction and HEIDI, and Bayesian coloc.abf supported a shared "
         "causal variant at the FDPS locus (PP.H4 = 0.9985), together "
         "indicating a protective effect of genetically determined FDPS "
         "expression on breast cancer risk [44-46]. The protective direction "
         "was replicated in ER+ (P = 1.54e-05, HEIDI P = 1.09e-02) and ER- "
         "(P = 2.78e-01) OncoArray sensitivity analyses [37], with ER+ "
         "reaching Bonferroni significance and ER- underpowered. These results "
         "are consistent with Mendelian-randomization evidence that statins "
         "reduce ER+ breast cancer risk [47] and with the established adjuvant "
         "benefit of nitrogen-containing bisphosphonates, which inhibit FDPS "
         "[48]. Multi-signal colocalisation was less conclusive (maximum "
         "PP.H4 = 0.44)."),
    "Single-cell localization (Section 3.6)":
        ("",
         "All single-cell, SMR/HEIDI and colocalisation pipelines are archived "
         "and reusable; follow-up work could refine the FDPS causal inference "
         "with independent fine-mapping resources, larger ER- GWAS and "
         "tissue-matched eQTLs."),
}

# new section inserted after the final body paragraph (4.6 follow-up)
NEW_SECTIONS = [
    ("4.7 Comparison with published lipid-metabolism signatures\n",
     "Six published lipid-metabolism breast cancer signatures (2021-2025) "
     "were identified by systematic search [7,49-53] (Supplementary Table S11). "
     "All used TCGA-based LASSO feature selection; two validated across "
     "multiple cohorts, none performed germline causal inference, and none "
     "combined CIBERSORT/ESTIMATE with external subtype mapping. To our "
     "knowledge, this is the first cholesterol-focused breast cancer study "
     "to combine cross-platform ER-classifier validation, externally mapped "
     "subtypes, immune deconvolution, single-cell localization and germline "
     "causal inference in one evidence chain."),
]

# appended references (after renumbering old 42-46 -> 43-47)
REFS_APPEND = [
    "[48] Ma L, Qian B, Peng C, Liu G, Shen H. A lipid metabolism-related "
    "gene signature predicts prognosis after tamoxifen treatment in ER+ "
    "breast cancer and reflects tumor microenvironment heterogeneity through "
    "single-cell analysis. BMC Med Genomics. 2025;18:123.",
    "[49] Zhao X, Yan L, Yang Z, et al. A novel signature incorporating genes "
    "related to lipid metabolism and immune for prognostic and functional "
    "prediction of breast cancer. Aging (Albany NY). 2024;16:8611-8629.",
    "[50] Shen L, Huang H, Li J, et al. Exploration of prognosis and "
    "immunometabolism landscapes in ER+ breast cancer based on a novel lipid "
    "metabolism-related signature. Front Immunol. 2023;14:1199465.",
    "[51] Gong M, Liu X, Yang W, et al. Identification of a lipid "
    "metabolism-associated gene signature predicting survival in breast "
    "cancer. Int J Gen Med. 2021;14:9503-9513.",
    "[52] Chang X, Xing P. Identification of a novel lipid metabolism-related "
    "gene signature within the tumour immune microenvironment for breast "
    "cancer. Lipids Health Dis. 2022;21:43.",
]

HPA_REF = ("[42] Uhlen M, Fagerberg L, Hallstrom BM, et al. Proteomics. "
           "Tissue-based map of the human proteome. Science. "
           "2015;347:1260419.")

FACKLER_REF = ("[43] Fackler MJ, Umbricht CB, Williams D, et al. "
               "Genome-wide methylation analysis identifies genes specific "
               "to breast cancer hormone receptor status and risk of "
               "recurrence. Cancer Res. 2011;71:6195-6207.")

NEWMAN_REF = ("[28] Newman AM, Liu CL, Green MR, et al. Robust enumeration of "
              "cell subsets from tissue expression profiles. Nat Methods. "
              "2015;12:453-457.")

YOSHIHARA_REF = ("[29] Yoshihara K, Shahmoradgoli M, Martinez E, et al. "
                  "Inferring tumour purity and stromal and immune cell "
                  "admixture from expression data. Bioinformatics. "
                  "2013;29:2112-2120.")

# old figure number -> new figure number (first-mention order 1..16)
FIG_MAP = {
    1: 1, 2: 2, 7: 3, 3: 4, 4: 5, 5: 6, 8: 7, 14: 8, 6: 9,
    13: 10, 9: 11, "10b": 12, 10: 13, 11: 14, "12b": 15, 12: 16,
}


def fig_renumber(text):
    """Renumber in-text 'Figure N' / caption references."""
    def repl(m):
        old = m.group(1)
        key = old if old in FIG_MAP else int(old)
        return f"Figure {FIG_MAP[key]}"
    return re.sub(r"Figure\s+(\d+[a-z]?)\b", repl, text)


def renumber_caption(p):
    """Renumber a caption paragraph, preserving the bold 'Figure N.' prefix."""
    t = p.text
    if not t.startswith("Figure "):
        return False
    m = re.match(r"^(Figure\s+\d+[a-z]?\.)(.*)$", t)
    if not m:
        return False
    prefix, rest = m.group(1), m.group(2)
    new_prefix = fig_renumber(prefix)
    if new_prefix == prefix:
        return False
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    r = p.add_run(new_prefix + " ")
    r.bold = True
    p.add_run(rest.strip())
    return True


def set_text(p, bold_part, body_part):
    """Replace paragraph runs, preserving paragraph style."""
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if bold_part:
        r = p.add_run(bold_part.rstrip("\n"))
        r.bold = True
        if bold_part.endswith("\n"):
            p.add_run().add_break()
    lines = body_part.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        if line:
            p.add_run(line)


def reflow_figures(doc):
    """Move all (image, caption) blocks after the references, in order 1..16."""
    rel_map = {rid: r.target_ref for rid, r in doc.part.rels.items()
               if "image" in r.reltype}
    media_to_fig = {
        "image1.png": 1, "image2.png": 2, "image3.png": 4, "image4.png": 5,
        "image5.png": 6, "image6.png": 9, "image7.png": 3, "image8.png": 7,
        "image9.png": 11, "image10.png": 13, "image11.png": 14,
        "image12.png": 16, "image13.png": 8, "image14.png": 10,
        "image15.png": 12, "image16.png": 15,
    }
    img_paras, cap_paras = {}, {}
    for p in doc.paragraphs:
        blips = p._p.findall(".//" + qn("a:blip"))
        if not blips:
            continue
        rid = blips[0].get(qn("r:embed"))
        base = (rel_map.get(rid, "") or "").rsplit("/", 1)[-1]
        if base in media_to_fig:
            img_paras[media_to_fig[base]] = p
    for p in doc.paragraphs:
        m = re.match(r"^Figure (\d+)\.", p.text.strip())
        if m:
            cap_paras[int(m.group(1))] = p
    assert set(img_paras) == set(range(1, 17)), sorted(img_paras)
    assert set(cap_paras) == set(range(1, 17)), sorted(cap_paras)

    body = doc.element.body
    for n in range(1, 17):
        for el in (img_paras[n]._p, cap_paras[n]._p):
            el.getparent().remove(el)

    heading = body.makeelement(qn("w:p"), {})
    body.append(heading)
    hp = Paragraph(heading, doc)
    r = hp.add_run("Figure Legends")
    r.bold = True
    for n in range(1, 17):
        body.append(img_paras[n]._p)
        body.append(cap_paras[n]._p)
    return img_paras, cap_paras


def main():
    doc = docx.Document(V15)
    paras = [x for x in doc.paragraphs if x.text.strip()]
    by_prefix = {}
    for p in paras:
        key = p.text.strip()[:40]
        by_prefix.setdefault(key, []).append(p)

    # 0) figure renumbering first, so inserted replacement text is not
    #    re-mapped a second time.
    for p in paras:
        if renumber_caption(p):
            continue
        t = p.text
        if "Figure " in t:
            new = fig_renumber(t)
            if new != t:
                set_text(p, "", new)

    # 1) abstract
    abs_idx = next(i for i, p in enumerate(paras) if p.text.strip() == "Abstract")
    for i, (label, body) in enumerate(ABSTRACT, start=1):
        set_text(paras[abs_idx + i], label, body)

    # 2) keywords
    kw = next(p for p in paras if p.text.strip().startswith("Keywords:"))
    set_text(kw, "", KEYWORDS)

    # 3) body replacements
    matched = set()
    for key, (bold, body) in BODY.items():
        hits = [p for p in paras if p.text.strip().startswith(key)]
        assert len(hits) == 1, f"{key!r}: {len(hits)} matches"
        set_text(hits[0], bold, body)
        matched.add(hits[0]._p)

    # 3.5) insert new section 4.7 after the final body paragraph
    anchor = next(p for p in paras
                  if p.text.strip().startswith("All single-cell, SMR/HEIDI"))
    for bold, body in NEW_SECTIONS:
        el = doc.element.body.makeelement(qn("w:p"), {})
        anchor._p.addnext(el)
        np_ = Paragraph(el, doc)
        set_text(np_, bold, body)
        anchor = np_

    # 3.6) references: insert CIBERSORT/ESTIMATE as [28,29], renumber
    #      28-39 -> 30-41, drop old 40/41, renumber 42-46 -> 43-47,
    #      insert HPA as [42], append [48]-[52], then insert Fackler as
    #      [43] and shift 43-52 -> 44-53.
    refs = [p for p in paras if re.match(r"^\[\d+\]", p.text.strip())]
    p28_old = next(p for p in refs if p.text.strip().startswith("[28] "))
    p42_old = next(p for p in refs if p.text.strip().startswith("[42] "))
    p46_old = next(p for p in refs if p.text.strip().startswith("[46] "))
    p40_old = next(p for p in refs if p.text.strip().startswith("[40] "))
    p41_old = next(p for p in refs if p.text.strip().startswith("[41] "))
    for p in (p40_old, p41_old):
        p._p.getparent().remove(p._p)
    for old, new in ((39, 41), (38, 40), (37, 39), (36, 38), (35, 37),
                     (34, 36), (33, 35), (32, 34), (31, 33), (30, 32),
                     (29, 31), (28, 30)):
        tgt = next(p for p in refs if p.text.strip().startswith(f"[{old}] "))
        rest = tgt.text.strip()[len(f"[{old}] "):]
        set_text(tgt, "", f"[{new}] {rest}")
    for old, new in ((46, 47), (45, 46), (44, 45), (43, 44), (42, 43)):
        tgt = next(p for p in refs if p.text.strip().startswith(f"[{old}] "))
        rest = tgt.text.strip()[len(f"[{old}] "):]
        set_text(tgt, "", f"[{new}] {rest}")
    # insert Newman [28] and Yoshihara [29] before old [28] (now [30])
    el28 = doc.element.body.makeelement(qn("w:p"), {})
    p28_old._p.addprevious(el28)
    set_text(Paragraph(el28, doc), "", NEWMAN_REF)
    el29 = doc.element.body.makeelement(qn("w:p"), {})
    el28.addnext(el29)
    set_text(Paragraph(el29, doc), "", YOSHIHARA_REF)
    # insert HPA as [42] before old [42] (now [43])
    el44 = doc.element.body.makeelement(qn("w:p"), {})
    p42_old._p.addprevious(el44)
    set_text(Paragraph(el44, doc), "", HPA_REF)
    # append [48]-[52] after old [46] (now [47])
    cursor = p46_old
    for ref in REFS_APPEND:
        el = doc.element.body.makeelement(qn("w:p"), {})
        cursor._p.addnext(el)
        np_ = Paragraph(el, doc)
        set_text(np_, "", ref)
        cursor = np_

    # insert Fackler as [43] and shift current 43-52 -> 44-53 (reverse)
    live = [p for p in doc.paragraphs if re.match(r"^\[\d+\]", p.text.strip())]
    for old, new in ((52, 53), (51, 52), (50, 51), (49, 50), (48, 49),
                     (47, 48), (46, 47), (45, 46), (44, 45), (43, 44)):
        tgt = next(p for p in live if p.text.strip().startswith(f"[{old}] "))
        rest = tgt.text.strip()[len(f"[{old}] "):]
        set_text(tgt, "", f"[{new}] {rest}")
    el43 = doc.element.body.makeelement(qn("w:p"), {})
    p42_old._p.addprevious(el43)
    set_text(Paragraph(el43, doc), "", FACKLER_REF)

    # 4) declarations
    da = next(p for p in paras if p.text.strip().startswith("Data availability."))
    ca = next(p for p in paras if p.text.strip().startswith("Code availability."))
    set_text(da, "Data availability. ",
             "All data are public: TCGA-BRCA (NCI GDC; accessed 4-6 August "
             "2026) and TCGA-CDR [14]; METABRIC (cBioPortal) [15]; PAM50 calls "
             "(UCSC Xena); FDPS HM450 methylation (cBioPortal); GSE21653 (NCBI "
             "GEO); GSE7390 (ArrayExpress E-GEOD-7390); GSE176078 (CELLxGENE); "
             "GSE20711 (NCBI GEO); "
             "GSE161529 (Mendeley Data mirror of the Pal et al. atlas [12]); "
             "eQTLGen [35]; GTEx v8 [38,39]; BCAC GWAS (GWAS Catalog "
             "GCST010098/GCST010100/GCST004988; Zhang et al. 2020 [36] and "
             "Michailidou et al. 2017 [37]); 1000 Genomes Phase 3 EUR. Curated "
             "intermediate results are provided in the GitHub repository "
             "(" + GITHUB_URL + ") and the Zenodo archive (" + ZENODO_DOI +
             "); raw data must be obtained from the original sources under "
             "their respective data-use policies. Supplementary Tables S1-S15 "
             "and Supplementary Figures S1-S3 accompany this manuscript; the "
             "TRIPOD-AI checklist (Supplementary Table S12), a comparison "
             "with published lipid-metabolism signatures (S11) and a "
             "graphical abstract are also provided in the repository.")
    set_text(ca, "Code availability. ",
             "All analysis scripts (bc_01..bc_07, validation, single-cell, "
             "SMR/HEIDI, colocalisation, immune-deconvolution and bonus "
             "pipelines) and curated results are "
             "publicly available at GitHub (" + GITHUB_URL + ") and archived "
             "in Zenodo (" + ZENODO_DOI + ").")

    # 5) reflow figure blocks after the references, in numerical order
    reflow_figures(doc)

    doc.save(V20)
    print("saved", V20)


if __name__ == "__main__":
    main()
