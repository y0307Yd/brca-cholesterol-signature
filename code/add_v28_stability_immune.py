from __future__ import annotations

import json
import os
import re
from pathlib import Path

import docx
import matplotlib as mpl
import numpy as np
import pandas as pd
from docx.shared import Inches
from scipy import stats
from scipy.cluster.hierarchy import fcluster, linkage
from scipy.optimize import linear_sum_assignment
from scipy.spatial.distance import squareform
from sklearn.cluster import KMeans
from sklearn.metrics import adjusted_rand_score, silhouette_score

os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parent / ".mplconfig"))
mpl.use("Agg")
import matplotlib.pyplot as plt


SEED = 20260804
REPO = Path(__file__).resolve().parents[1]
SOURCE = Path(os.environ.get("BRCA_CHOL_PROJECT_ROOT", REPO)).resolve()
DATA = SOURCE / "outputs" / "chol_metab_signature"
OUT = Path(os.environ.get("BRCA_CHOL_OUTPUT_DIR", REPO / "reproduction_output")).resolve()
OUT.mkdir(exist_ok=True)

mpl.rcParams.update({
    "font.family": "sans-serif",
    "font.sans-serif": ["Arial", "Helvetica", "DejaVu Sans", "sans-serif"],
    "font.size": 7,
    "axes.spines.right": False,
    "axes.spines.top": False,
    "axes.linewidth": 0.8,
    "legend.frameon": False,
    "pdf.fonttype": 42,
    "svg.fonttype": "none",
})

FAMILY_COLORS = {
    "ESTIMATE": "#2b6f8a",
    "Marker score": "#4f9d8a",
    "ssGSEA": "#7b6fa6",
    "CIBERSORT": "#c58a3a",
    "Checkpoint": "#b44b4b",
}


def bh_adjust(values: pd.Series) -> np.ndarray:
    p = values.to_numpy(float)
    order = np.argsort(p)
    ranked = p[order]
    adjusted = ranked * len(p) / np.arange(1, len(p) + 1)
    adjusted = np.minimum.accumulate(adjusted[::-1])[::-1]
    out = np.empty_like(adjusted)
    out[order] = np.minimum(adjusted, 1.0)
    return out


def save_figure(fig: plt.Figure, stem: str) -> None:
    base = OUT / stem
    fig.savefig(base.with_suffix(".png"), dpi=600, bbox_inches="tight")
    fig.savefig(base.with_suffix(".pdf"), bbox_inches="tight")
    fig.savefig(base.with_suffix(".svg"), bbox_inches="tight")
    fig.savefig(base.with_suffix(".tiff"), dpi=600, bbox_inches="tight")


def load_hub_z() -> tuple[np.ndarray, list[str]]:
    genes = [x.strip() for x in open(DATA / "tcga_genes.txt", encoding="utf-8") if x.strip()]
    hubs = pd.read_csv(DATA / "hub_genes.csv")["gene"].tolist()
    idx = {g: i for i, g in enumerate(genes)}
    x = np.load(DATA / "tcga_Xlog.npy", mmap_mode="r")
    z = np.vstack([(x[idx[g]] - x[idx[g]].mean()) / x[idx[g]].std() for g in hubs]).T
    return np.asarray(z, float), hubs


def labels_from_consensus(k: int) -> tuple[np.ndarray, np.ndarray]:
    matrix = np.load(DATA / f"consensus_k{k}.npy")
    labels = fcluster(linkage(squareform(1 - matrix), method="average"),
                      t=k, criterion="maxclust")
    return labels, matrix


def map_labels(reference: np.ndarray, candidate: np.ndarray, k: int) -> np.ndarray:
    contingency = np.zeros((k, k), dtype=int)
    for i in range(1, k + 1):
        for j in range(k):
            contingency[i - 1, j] = np.sum((reference == i) & (candidate == j))
    rows, cols = linear_sum_assignment(-contingency)
    mapping = {col: row + 1 for row, col in zip(rows, cols)}
    return np.asarray([mapping[x] for x in candidate], int)


def subtype_stability(n_resamples: int = 200) -> tuple[pd.DataFrame, pd.DataFrame]:
    z, _ = load_hub_z()
    n = len(z)
    rng = np.random.default_rng(SEED + 28)
    summary_rows, resample_rows = [], []
    for k in (3, 4, 5):
        locked, consensus = labels_from_consensus(k)
        off = consensus[np.triu_indices(n, 1)]
        pac = float(np.mean((off > 0.1) & (off < 0.9)))
        cluster_consensus = []
        for c in range(1, k + 1):
            ix = np.where(locked == c)[0]
            cluster_consensus.append(float(consensus[np.ix_(ix, ix)][np.triu_indices(len(ix), 1)].mean()))
        sil = float(silhouette_score(z, locked, metric="euclidean"))

        for b in range(n_resamples):
            sampled = rng.choice(n, size=int(0.8 * n), replace=False)
            model = KMeans(n_clusters=k, n_init=20, random_state=SEED + 1000 * k + b)
            model.fit(z[sampled])
            # Refit mapping on sampled observations, then apply it to all observations.
            contingency = np.zeros((k, k), dtype=int)
            raw_sampled = model.predict(z[sampled])
            for i in range(1, k + 1):
                for j in range(k):
                    contingency[i - 1, j] = np.sum((locked[sampled] == i) & (raw_sampled == j))
            rows, cols = linear_sum_assignment(-contingency)
            mapping = {col: row + 1 for row, col in zip(rows, cols)}
            predicted = np.asarray([mapping[x] for x in model.predict(z)], int)
            held_out = np.setdiff1d(np.arange(n), sampled)
            jaccards = []
            heldout_jaccards = []
            for c in range(1, k + 1):
                a, bset = locked == c, predicted == c
                jaccards.append(float(np.sum(a & bset) / np.sum(a | bset)))
                ah, bh = locked[held_out] == c, predicted[held_out] == c
                heldout_jaccards.append(float(np.sum(ah & bh) / np.sum(ah | bh)))
            resample_rows.append({
                "k": k,
                "resample": b + 1,
                "adjusted_rand_index": float(adjusted_rand_score(locked, predicted)),
                "sample_retention": float(np.mean(locked == predicted)),
                "mean_cluster_jaccard": float(np.mean(jaccards)),
                "minimum_cluster_jaccard": float(np.min(jaccards)),
                "heldout_adjusted_rand_index": float(adjusted_rand_score(locked[held_out], predicted[held_out])),
                "heldout_sample_retention": float(np.mean(locked[held_out] == predicted[held_out])),
                "heldout_mean_cluster_jaccard": float(np.mean(heldout_jaccards)),
                **{f"cluster_{c}_jaccard": jaccards[c - 1] for c in range(1, k + 1)},
            })

        r = pd.DataFrame([x for x in resample_rows if x["k"] == k])
        summary_rows.append({
            "k": k,
            "PAC_0.1_0.9": pac,
            "silhouette": sil,
            "mean_cluster_consensus": float(np.mean(cluster_consensus)),
            "minimum_cluster_consensus": float(np.min(cluster_consensus)),
            "median_adjusted_rand_index": float(r.adjusted_rand_index.median()),
            "ARI_2.5pct": float(r.adjusted_rand_index.quantile(0.025)),
            "ARI_97.5pct": float(r.adjusted_rand_index.quantile(0.975)),
            "median_sample_retention": float(r.sample_retention.median()),
            "retention_2.5pct": float(r.sample_retention.quantile(0.025)),
            "retention_97.5pct": float(r.sample_retention.quantile(0.975)),
            "median_mean_cluster_jaccard": float(r.mean_cluster_jaccard.median()),
            "median_minimum_cluster_jaccard": float(r.minimum_cluster_jaccard.median()),
            "median_heldout_adjusted_rand_index": float(r.heldout_adjusted_rand_index.median()),
            "median_heldout_sample_retention": float(r.heldout_sample_retention.median()),
            "heldout_retention_2.5pct": float(r.heldout_sample_retention.quantile(0.025)),
            "heldout_retention_97.5pct": float(r.heldout_sample_retention.quantile(0.975)),
            "median_heldout_mean_cluster_jaccard": float(r.heldout_mean_cluster_jaccard.median()),
            "n_resamples": n_resamples,
            "resample_fraction": 0.8,
        })
    summary = pd.DataFrame(summary_rows)
    resamples = pd.DataFrame(resample_rows)
    summary.to_csv(OUT / "Supplementary_Table_S26_subtype_stability.csv", index=False,
                   encoding="utf-8-sig")
    resamples.to_csv(OUT / "Supplementary_Table_S26_subtype_stability_resamples.csv", index=False,
                     encoding="utf-8-sig")
    return summary, resamples


def plot_stability(summary: pd.DataFrame, resamples: pd.DataFrame) -> None:
    colors = {3: "#8d99a6", 4: "#2b6f8a", 5: "#c58a3a"}
    fig, axes = plt.subplots(2, 2, figsize=(7.2, 5.7))
    ax = axes[0, 0]
    ax.bar(summary.k.astype(str), summary["PAC_0.1_0.9"],
           color=[colors[k] for k in summary.k])
    ax.set_ylabel("PAC (lower is better)")
    ax.set_title("a  Ambiguous clustering", loc="left", fontweight="bold", fontsize=8)
    ax.grid(axis="y", alpha=.2)

    ax = axes[0, 1]
    width = .35
    x = np.arange(3)
    ax.bar(x - width / 2, summary.silhouette, width, color="#4f9d8a", label="Silhouette")
    ax.bar(x + width / 2, summary.mean_cluster_consensus, width, color="#7b6fa6",
           label="Mean cluster consensus")
    ax.set_xticks(x, [f"k={k}" for k in summary.k])
    ax.set_ylim(0, 1)
    ax.set_title("b  Separation and within-cluster consensus", loc="left",
                 fontweight="bold", fontsize=8)
    ax.legend(fontsize=6, loc="upper left")
    ax.grid(axis="y", alpha=.2)

    for ax, metric, title, ylabel in [
        (axes[1, 0], "adjusted_rand_index", "c  Global label reproducibility", "Adjusted Rand index"),
        (axes[1, 1], "mean_cluster_jaccard", "d  Cluster membership reproducibility", "Mean cluster Jaccard"),
    ]:
        groups = [resamples.loc[resamples.k == k, metric].to_numpy() for k in (3, 4, 5)]
        parts = ax.violinplot(groups, positions=[1, 2, 3], widths=.75, showmedians=True,
                              showextrema=False)
        for body, k in zip(parts["bodies"], (3, 4, 5)):
            body.set_facecolor(colors[k]); body.set_alpha(.8)
        ax.set_xticks([1, 2, 3], ["k=3", "k=4", "k=5"])
        ax.set_ylim(0, 1)
        ax.set_ylabel(ylabel)
        ax.set_title(title, loc="left", fontweight="bold", fontsize=8)
        ax.grid(axis="y", alpha=.2)
    fig.suptitle("Resampling stability of the hub-gene molecular subtypes", fontsize=10, y=.99)
    fig.tight_layout(rect=(0, 0, 1, .97))
    save_figure(fig, "Supplementary_Figure_S13_subtype_stability")
    plt.close(fig)


def parse_estimate() -> pd.DataFrame:
    tab = pd.read_csv(SOURCE / "work" / "estimate_scores.txt", sep="\t", skiprows=2, index_col=0)
    tab = tab.drop(columns=[c for c in tab.columns if c == "Description"], errors="ignore")
    tab.columns = [str(c).replace(".", "-") for c in tab.columns]
    tab = tab.T
    tab.index = [re.sub(r"-(01[AB]).*$", r"-\1", x) for x in tab.index]
    return tab.apply(pd.to_numeric, errors="coerce")


def parse_pam50() -> pd.Series:
    df = pd.read_csv(SOURCE / "data" / "pam50" / "BRCA_clinicalMatrix.tsv",
                     sep="\t", dtype=str, low_memory=False)
    s = df.set_index("sampleID")["PAM50Call_RNAseq"].dropna()
    s.index = ["-".join(str(x).split("-")[:3]) for x in s.index]
    return s[~s.index.duplicated(keep="first")]


def build_immune_data() -> pd.DataFrame:
    clinical = pd.read_csv(DATA / "tcga_subtypes.csv")
    clinical["patient_key"] = clinical.sample_id.str.split("-").str[:3].str.join("-")
    clinical["PAM50"] = clinical.patient_key.map(parse_pam50())
    est = parse_estimate().reindex(clinical.sample_id)
    if est.isna().all(axis=1).any():
        raise RuntimeError("ESTIMATE sample mapping is incomplete")

    records = []
    for feature in est.columns:
        records.append(pd.DataFrame({"sample_id": clinical.sample_id, "family": "ESTIMATE",
                                     "feature": feature, "value": est[feature].to_numpy()}))

    markers = pd.read_csv(DATA / "immune_scores.csv")
    if not np.array_equal(markers.subtype, clinical.subtype):
        raise RuntimeError("Marker-score row order does not match subtype labels")
    for feature in markers.columns.drop("subtype"):
        records.append(pd.DataFrame({"sample_id": clinical.sample_id, "family": "Marker score",
                                     "feature": feature, "value": markers[feature]}))

    ssgsea = pd.read_csv(DATA / "ssgsea_scores.csv")
    expected = clinical.sample_id.str.replace(r"-(01)[AB]$", r"-\1", regex=True)
    if (not expected.is_unique or not ssgsea.sample_norm.is_unique or
            not np.array_equal(ssgsea.sample_norm, expected) or
            not np.array_equal(ssgsea.subtype, clinical.subtype)):
        raise RuntimeError("ssGSEA sample mapping does not match TCGA order")
    for feature in ssgsea.columns.drop(["sample_norm", "subtype"]):
        records.append(pd.DataFrame({"sample_id": clinical.sample_id, "family": "ssGSEA",
                                     "feature": feature, "value": ssgsea[feature]}))

    cib = pd.read_csv(SOURCE / "work" / "bonus_cibersort_fractions.csv")
    if not np.array_equal(cib.sample_id, clinical.sample_id):
        raise RuntimeError("CIBERSORT sample order does not match TCGA order")
    for feature in cib.columns.drop(["sample_id", "RMSE"]):
        value = np.arcsin(np.sqrt(np.clip(cib[feature].to_numpy(float), 0, 1)))
        records.append(pd.DataFrame({"sample_id": clinical.sample_id, "family": "CIBERSORT",
                                     "feature": feature, "value": value}))

    checkpoint_genes = ["CD274", "PDCD1", "CTLA4", "LAG3", "HAVCR2", "IDO1", "CD8A"]
    genes = [x.strip() for x in open(DATA / "tcga_genes.txt", encoding="utf-8") if x.strip()]
    gene_index = {g: i for i, g in enumerate(genes)}
    x = np.load(DATA / "tcga_Xlog.npy", mmap_mode="r")
    for feature in checkpoint_genes:
        records.append(pd.DataFrame({"sample_id": clinical.sample_id, "family": "Checkpoint",
                                     "feature": feature, "value": np.asarray(x[gene_index[feature]], float)}))
    long = pd.concat(records, ignore_index=True)
    return long.merge(clinical[["sample_id", "subtype", "ER", "stage", "PAM50"]],
                      on="sample_id", validate="many_to_one").merge(
        est[["ESTIMATEScore"]].reset_index(names="sample_id"), on="sample_id", validate="many_to_one")


def design_matrix(df: pd.DataFrame, model: str, include_subtype: bool,
                  include_estimate: bool) -> np.ndarray:
    cols = [np.ones(len(df)), pd.to_numeric(df.stage).to_numpy(float)]
    if include_estimate:
        cols.append(pd.to_numeric(df.ESTIMATEScore).to_numpy(float))
    if model == "ER":
        cols.append(pd.to_numeric(df.ER).to_numpy(float))
    else:
        d = pd.get_dummies(df.PAM50, drop_first=True, dtype=float)
        cols.extend([d[c].to_numpy() for c in d.columns])
    if include_subtype:
        d = pd.get_dummies(df.subtype.astype(str), drop_first=True, dtype=float)
        cols.extend([d[c].to_numpy() for c in d.columns])
    return np.column_stack(cols)


def nested_test(y: np.ndarray, x0: np.ndarray, x1: np.ndarray) -> tuple[float, float, float]:
    rss0 = float(np.sum((y - x0 @ np.linalg.lstsq(x0, y, rcond=None)[0]) ** 2))
    rss1 = float(np.sum((y - x1 @ np.linalg.lstsq(x1, y, rcond=None)[0]) ** 2))
    q = x1.shape[1] - x0.shape[1]
    df2 = len(y) - x1.shape[1]
    f_stat = max(0.0, ((rss0 - rss1) / q) / (rss1 / df2))
    return f_stat, float(stats.f.sf(f_stat, q, df2)), float(max(0, (rss0 - rss1) / rss0))


def immune_models() -> pd.DataFrame:
    data = build_immune_data()
    rows = []
    for (family, feature), group in data.groupby(["family", "feature"], sort=False):
        groups = [group.loc[group.subtype == c, "value"].dropna().to_numpy() for c in range(1, 5)]
        kw = float(stats.kruskal(*groups).pvalue)
        include_estimate = family != "ESTIMATE"
        for model in ("ER", "PAM50"):
            cols = ["value", "subtype", "stage", model]
            if include_estimate:
                cols.append("ESTIMATEScore")
            complete = group[cols].dropna().copy()
            x0 = design_matrix(complete, model, False, include_estimate)
            x1 = design_matrix(complete, model, True, include_estimate)
            f_stat, p, partial_r2 = nested_test(complete.value.to_numpy(float), x0, x1)
            rows.append({
                "family": family,
                "feature": feature,
                "transformation": "arcsine_sqrt" if family == "CIBERSORT" else "none",
                "model": model,
                "covariates": ("stage + " + model + (" + ESTIMATEScore" if include_estimate else "")),
                "n": len(complete),
                "unadjusted_Kruskal_P": kw,
                "global_subtype_F": f_stat,
                "global_subtype_P": p,
                "partial_R2_subtype": partial_r2,
            })
    result = pd.DataFrame(rows)
    result["global_subtype_FDR_all_features"] = np.nan
    result["global_subtype_FDR_within_family"] = np.nan
    for model in result.model.unique():
        ix = result.model == model
        result.loc[ix, "global_subtype_FDR_all_features"] = bh_adjust(result.loc[ix, "global_subtype_P"])
        for family in result.family.unique():
            jx = ix & (result.family == family)
            result.loc[jx, "global_subtype_FDR_within_family"] = bh_adjust(result.loc[jx, "global_subtype_P"])
    result.to_csv(OUT / "Supplementary_Table_S27_adjusted_immune_models.csv", index=False,
                  encoding="utf-8-sig")
    return result


def plot_immune(result: pd.DataFrame) -> None:
    family_order = ["ESTIMATE", "Marker score", "ssGSEA", "CIBERSORT", "Checkpoint"]
    model_colors = {"ER": "#2b6f8a", "PAM50": "#b44b4b"}
    fig = plt.figure(figsize=(7.2, 10.2))
    gs = fig.add_gridspec(2, 1, height_ratios=[1, 5.4], hspace=.28)
    ax = fig.add_subplot(gs[0])
    counts = (result.assign(sig=result.global_subtype_FDR_all_features < .05)
              .groupby(["family", "model"], observed=True).sig.sum().unstack(fill_value=0)
              .reindex(family_order))
    x = np.arange(len(family_order)); width = .34
    ax.bar(x - width / 2, counts["ER"], width, color=model_colors["ER"], label="ER-adjusted")
    ax.bar(x + width / 2, counts["PAM50"], width, color=model_colors["PAM50"], label="PAM50-adjusted")
    totals = result[result.model == "ER"].groupby("family").size().reindex(family_order)
    ax.set_xticks(x, [f"{f}\n(n={totals[f]})" for f in family_order])
    ax.set_ylabel("Features at global FDR < 0.05")
    ax.set_title("a  Independent subtype signal by immune-feature family", loc="left",
                 fontweight="bold", fontsize=8)
    ax.legend(ncol=2, loc="upper right")
    ax.grid(axis="y", alpha=.2)

    ax = fig.add_subplot(gs[1])
    er = result[result.model == "ER"].copy()
    er["family"] = pd.Categorical(er.family, family_order, ordered=True)
    er = er.sort_values(["family", "partial_R2_subtype"], ascending=[True, False]).reset_index(drop=True)
    labels = [f"{r.feature}  ({r.family})" for r in er.itertuples()]
    y = np.arange(len(er))[::-1]
    pam = result[result.model == "PAM50"].set_index(["family", "feature"]).loc[
        list(zip(er.family.astype(str), er.feature))].reset_index()
    for model, frame, offset in [("ER", er, .13), ("PAM50", pam, -.13)]:
        sig = frame.global_subtype_FDR_all_features.to_numpy() < .05
        ax.scatter(frame.partial_R2_subtype, y + offset, s=np.where(sig, 25, 13),
                   facecolors=model_colors[model] if model == "ER" else "white",
                   edgecolors=model_colors[model], linewidths=.8,
                   label=f"{model}-adjusted" + (" (filled)" if model == "ER" else " (open)"))
    for i in range(1, len(er)):
        if er.family.iloc[i] != er.family.iloc[i - 1]:
            ax.axhline(len(er) - i - .5, color="#d5d9dc", lw=.7)
    ax.axvline(0, color="#777777", lw=.7)
    ax.set_yticks(y, labels, fontsize=5.3)
    ax.set_xlabel("Incremental variance explained by subtype (partial R2)")
    ax.set_title("b  Adjusted global subtype effects across all 50 immune features", loc="left",
                 fontweight="bold", fontsize=8)
    ax.legend(ncol=2, loc="lower right", fontsize=6)
    ax.grid(axis="x", alpha=.2)
    fig.suptitle("Immune phenotype after adjustment for ER status or PAM50, stage and tumour context",
                 fontsize=10, y=.995)
    fig.tight_layout(rect=(0, 0, 1, .985))
    save_figure(fig, "Supplementary_Figure_S14_adjusted_immune_phenotypes")
    plt.close(fig)


def replace_paragraph(paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.add_run(text)


def patch_docx(stability: pd.DataFrame, immune: pd.DataFrame) -> Path:
    src = OUT / "Manuscript_Cholesterol_Metabolism_BRCA_v27.docx"
    dst = OUT / "Manuscript_Cholesterol_Metabolism_BRCA_v28.docx"
    doc = docx.Document(src)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    k4 = stability.set_index("k").loc[4]
    p26b = next(p for p in paras if p.text.strip().startswith("Consensus clustering [26]"))
    p26b.add_run(
        " Post hoc subtype-stability analysis compared k = 3, 4 and 5 using silhouette width and within-cluster "
        "consensus from the locked consensus matrices. Reproducibility was assessed in 200 independent 80% "
        "subsamples without replacement. K-means labels were aligned to the locked consensus labels by optimal "
        "assignment; adjusted Rand index (ARI), overall sample-label retention and cluster-wise Jaccard similarity "
        "were then calculated after prediction of all samples from each resampled fit."
    )
    p26b.add_run(
        " For the immune sensitivity analysis, all ESTIMATE scores, nine marker scores, nine ssGSEA scores, 22 "
        "CIBERSORT fractions and seven checkpoint genes were tested using nested linear models comparing covariates "
        "alone with covariates plus subtype. Separate models adjusted for ER status or PAM50 and stage; non-ESTIMATE "
        "outcomes additionally included ESTIMATEScore, whereas ESTIMATE outcomes did not. CIBERSORT fractions were "
        "arcsine-square-root transformed. BH-FDR was controlled across all 50 features within each model, and subtype "
        "partial R2 quantified incremental variance explained."
    )

    p34 = next(p for p in paras if p.text.strip().startswith("Consensus clustering of hub-gene expression"))
    p34.add_run(
        f" In the 200-resample stability analysis, k = 4 showed median ARI {k4.median_adjusted_rand_index:.3f} "
        f"(95% resampling interval {k4['ARI_2.5pct']:.3f}-{k4['ARI_97.5pct']:.3f}), median sample-label "
        f"retention {k4.median_sample_retention:.1%}, held-out retention "
        f"{k4.median_heldout_sample_retention:.1%} and median mean cluster Jaccard "
        f"{k4.median_mean_cluster_jaccard:.3f}. Its silhouette width was {k4.silhouette:.3f} and mean within-cluster "
        f"consensus was {k4.mean_cluster_consensus:.3f}. The four-cluster solution was therefore reproducible and "
        f"had the lowest PAC, although k = 3 was a more parsimonious competitor with slightly higher silhouette "
        f"and resampling indices; k = 4 should not be viewed as uniquely determined (Supplementary Figure S13; "
        f"Supplementary Table S26)."
    )

    counts = (immune.assign(sig=immune.global_subtype_FDR_all_features < .05)
              .groupby(["model", "family"]).sig.sum())
    er_total = int((immune.query("model == 'ER'").global_subtype_FDR_all_features < .05).sum())
    pam_total = int((immune.query("model == 'PAM50'").global_subtype_FDR_all_features < .05).sum())
    strongest = (immune.query("model == 'PAM50'")
                 .sort_values("partial_R2_subtype", ascending=False).iloc[0])
    p35 = next(p for p in paras if p.text.strip().startswith("3.5 Immune microenvironment"))
    p35.add_run(
        f" In the unified multivariable analysis of all 50 immune features, subtype remained significant after "
        f"global FDR correction for {er_total}/50 features in the ER-adjusted models and {pam_total}/50 in the "
        f"PAM50-adjusted models. The corresponding PAM50-adjusted counts were "
        f"{int(counts.get(('PAM50', 'ESTIMATE'), 0))}/3 ESTIMATE, "
        f"{int(counts.get(('PAM50', 'Marker score'), 0))}/9 marker-score, "
        f"{int(counts.get(('PAM50', 'ssGSEA'), 0))}/9 ssGSEA, "
        f"{int(counts.get(('PAM50', 'CIBERSORT'), 0))}/22 CIBERSORT and "
        f"{int(counts.get(('PAM50', 'Checkpoint'), 0))}/7 checkpoint features. The largest PAM50-adjusted "
        f"incremental effect was {strongest.feature} ({strongest.family}; partial R2 = "
        f"{strongest.partial_R2_subtype:.3f}, global FDR = {strongest.global_subtype_FDR_all_features:.2e}). "
        f"Thus, ER/PAM50 and tumour context explain part, but not all, of the subtype-associated immune variation "
        f"(Supplementary Figure S14; Supplementary Table S27)."
    )

    p43 = next(p for p in paras if p.text.strip().startswith("4.3 Subtypes, proliferation and immunity"))
    p43.add_run(
        f" The broader adjusted analysis retained {pam_total}/50 immune features after PAM50-based global FDR "
        "control, indicating that the immune phenotype is not solely a proxy for intrinsic subtype. Nevertheless, "
        "attenuation of several features and the observational design argue against interpreting these associations "
        "as independent treatment-predictive biomarkers."
    )
    pav = next(p for p in paras if p.text.strip().startswith("Data availability."))
    replace_paragraph(pav, pav.text.replace("Supplementary Tables S1-S25", "Supplementary Tables S1-S27")
                      .replace("Supplementary Figures S1-S12", "Supplementary Figures S1-S14"))

    # Supplementary legends are appended to retain the existing main-figure order and inline objects.
    doc.add_paragraph(
        "Supplementary Figure S13. Resampling stability of the hub-gene molecular subtypes. "
        "a, Proportion of ambiguous clustering (PAC; lower is better) for k = 3-5. b, Silhouette width and mean "
        "within-cluster consensus from the locked consensus matrices. c,d, Distributions of ARI and mean "
        "cluster-wise Jaccard similarity across 200 independent 80% subsamples. Each resample used k-means fitted "
        "to the sampled observations, optimal label alignment and prediction of all 952 TCGA-BRCA samples."
    )
    doc.add_paragraph(
        "Supplementary Figure S14. Immune phenotype after multivariable adjustment. a, Number of features with a "
        "global subtype effect at BH-FDR < 0.05 across all 50 features, stratified by feature family and adjustment "
        "model. b, Subtype partial R2 for all features in ER-adjusted and PAM50-adjusted nested linear models. "
        "Models included stage and, for non-ESTIMATE outcomes, ESTIMATEScore; CIBERSORT fractions were "
        "arcsine-square-root transformed. Filled and open points denote the two adjustment models; point size "
        "distinguishes features passing model-wide FDR < 0.05."
    )
    doc.save(dst)
    return dst


def write_legends() -> None:
    (OUT / "Supplementary_Figure_S13_legend.md").write_text(
        "# Supplementary Figure S13\n\nResampling stability of the hub-gene molecular subtypes. "
        "PAC, silhouette width, within-cluster consensus, adjusted Rand index and cluster-wise Jaccard "
        "similarity compare k = 3-5. Distributions summarize 200 independent 80% subsamples of 952 TCGA-BRCA "
        "samples.\n", encoding="utf-8")
    (OUT / "Supplementary_Figure_S14_legend.md").write_text(
        "# Supplementary Figure S14\n\nImmune phenotype after multivariable adjustment. Nested linear models tested "
        "the incremental global effect of subtype after ER or PAM50 and stage adjustment; non-ESTIMATE outcomes "
        "also included ESTIMATEScore. BH-FDR was controlled across all 50 features in each model.\n",
        encoding="utf-8")


def main() -> None:
    stability, resamples = subtype_stability()
    plot_stability(stability, resamples)
    immune = immune_models()
    plot_immune(immune)
    manuscript = patch_docx(stability, immune)
    write_legends()
    summary = {
        "stability_k4": stability.set_index("k").loc[4].to_dict(),
        "immune_significant_ER": int((immune.query("model == 'ER'").global_subtype_FDR_all_features < .05).sum()),
        "immune_significant_PAM50": int((immune.query("model == 'PAM50'").global_subtype_FDR_all_features < .05).sum()),
        "immune_features": int(immune[["family", "feature"]].drop_duplicates().shape[0]),
        "manuscript": str(manuscript),
    }
    (OUT / "v28_priority_additions_summary.json").write_text(json.dumps(summary, indent=2), encoding="utf-8")
    print(json.dumps(summary, indent=2))
    print(stability.to_string(index=False))
    print(immune.groupby(["model", "family"])["global_subtype_FDR_all_features"].apply(lambda x: int((x < .05).sum())))


if __name__ == "__main__":
    main()
