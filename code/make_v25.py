# -*- coding: utf-8 -*-
"""Generate v25 from v24: TCGA-vs-GTEx confirmation of normal-tumour axes."""
import docx

SRC = (r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature"
       r"\Manuscript_Cholesterol_Metabolism_BRCA_v24.docx")
OUT = (r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature"
       r"\Manuscript_Cholesterol_Metabolism_BRCA_v25.docx")


def set_text(p, bold_part, body_part):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if bold_part:
        r = p.add_run(bold_part.rstrip("\n"))
        r.bold = True
        if bold_part.endswith("\n"):
            p.add_run().add_break()
    lines = body_part.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        if line:
            p.add_run(line)


REPL = {
    "The four cholesterol-metabolism subtypes recapitulate": (
        "",
        "The four cholesterol-metabolism subtypes recapitulate known "
        "breast cancer biology: C1 is predominantly ER-negative, "
        "proliferative (cell-cycle up) and immune-infiltrated (highest "
        "CD8 T, regulatory T, NK, B-cell and macrophage scores), "
        "consistent with the immunologically hot, proliferative phenotype "
        "of triple-negative tumours; C3 is almost uniformly ER+, "
        "immune-cold and oxidative-phosphorylation-high; C2 shows "
        "down-regulated cell-cycle programs and high stromal/dendritic "
        "scores; C4 is a mixed, neutrophil-high group. Rank-based "
        "pathway-activity scores showed the strongest subtype differences "
        "for estrogen response (highest in the luminal-like C2/C3) and "
        "for cell-cycle and immune-cytokine programs (highest in C1), "
        "whereas cholesterol biosynthesis was only modestly higher in C4 "
        "than in C1 (Supplementary Figure S7). These patterns provide a "
        "cell-level interpretation of the signature even before "
        "single-cell data are added: the cholesterol-synthesis-low, "
        "ER-low C1 subtype is proliferative and immune-active [2,8], "
        "while cholesterol-synthesis activity per se does not segregate "
        "the luminal subtypes, indicating that the ER-classifier signal "
        "reflects ER-pathway biology more than cholesterol-pathway "
        "activity as a whole."),
    "Data availability.": (
        "Data availability. ",
        "All data are public: TCGA-BRCA (NCI GDC; accessed 4-6 August "
        "2026) and TCGA-CDR [14]; METABRIC (cBioPortal) [15]; PAM50 calls "
        "(UCSC Xena); FDPS HM450 methylation (cBioPortal); GSE21653 (NCBI "
        "GEO); GSE7390 (ArrayExpress E-GEOD-7390); GSE20711 (NCBI GEO); "
        "GSE15852 (NCBI GEO; paired normal-tumour); GSE91061 (NCBI GEO; "
        "immunotherapy RNA-seq); GSE176078 (CELLxGENE); GSE161529 "
        "(Mendeley Data mirror of the Pal et al. atlas [12]); eQTLGen "
        "[35]; GTEx v8 [38,39]; BCAC GWAS (GWAS Catalog "
        "GCST010098/GCST010100/GCST004988; Zhang et al. 2020 [36] and "
        "Michailidou et al. 2017 [37]); 1000 Genomes Phase 3 EUR. A "
        "complete data inventory with accessions and usage notes is "
        "provided in Supplementary Table S17. Curated intermediate "
        "results are provided in the GitHub repository "
        "(https://github.com/y0307Yd/brca-cholesterol-signature) and the "
        "Zenodo archive (https://doi.org/10.5281/zenodo.21873168); raw "
        "data must be obtained from the original sources under their "
        "respective data-use policies. Supplementary Tables S1-S20 and "
        "Supplementary Figures S1-S10 accompany this manuscript; the "
        "TRIPOD-AI checklist (Supplementary Table S12), a comparison with "
        "published lipid-metabolism signatures (S11) and a graphical "
        "abstract are also provided in the repository."),
    "2.1 Data sources and verification": (
        "2.1 Data sources and verification ",
        "The overall study design and the flow of analyses are summarised "
        "in Supplementary Figure S2. For validation analyses, PAM50 calls "
        "for TCGA-BRCA were obtained from UCSC Xena (GDC hub) [9]; FDPS "
        "HM450 methylation beta values were obtained from the cBioPortal "
        "data repository (brca_tcga_methylation_hm450); the independent "
        "expression cohorts GSE21653 (Affymetrix HG-U133 Plus 2.0, NCBI "
        "GEO with GPL570 annotation) [10] and GSE7390 (Affymetrix "
        "HG-U133A, ArrayExpress E-GEOD-7390 processed matrix and SDRF "
        "clinical data, probes mapped with hgu133a.db) [11] were used; a "
        "fourth independent Affymetrix cohort (GSE20711, NCBI GEO) was "
        "used for classifier transfer; the paired normal-tumour cohort "
        "GSE15852 (Affymetrix HG-U133A; 43 pairs) was used for "
        "normal-versus-tumour expression comparisons; the immunotherapy "
        "cohort GSE91061 (RNA-seq, melanoma anti-PD-1; 51 pre-treatment "
        "samples) was used for exploratory immunotherapy-response "
        "associations; and the second single-cell atlas GSE161529 "
        "(Mendeley Data mirror of the Pal et al. atlas) [12] was used "
        "for single-cell validation. The Affymetrix annotation was taken "
        "from the GEO GPL570 platform file. TCGA-BRCA RNA-sequencing "
        "STAR counts, methylation, mutation and clinical files were "
        "obtained from the NCI GDC portal [13] and verified against the "
        "official manifest by MD5 checksums (4,393/4,393 files); "
        "progression-free interval (PFI) was taken from TCGA-CDR [14]. "
        "METABRIC expression and clinical annotation with recurrence-free "
        "survival (RFS) were obtained from cBioPortal [15]. TCGA ER/PR/"
        "HER2 status was parsed from clinical XML; METABRIC receptor "
        "status from clinical sample files."),
}

APPEND = []

INSERT_AFTER = []

# v25 overrides: TCGA-vs-GTEx confirmation added to section 3.16
REPL["3.16 Cholesterol genes are remodelled in tumour versus normal "
     "breast tissue"] = (
    "3.16 Cholesterol genes are remodelled in tumour versus normal "
    "breast tissue\n",
    "In a paired normal-tumour cohort (GSE15852, 43 pairs), cholesterol "
    "biosynthetic enzymes were significantly up-regulated in tumours: "
    "DHCR24 (log2 fold change 0.59, paired Wilcoxon P = 7e-4), DHCR7 "
    "(0.82, P = 2.8e-3), HSD17B7 (0.36, P = 1.2e-3), HMGCS2 (0.83, "
    "P = 0.13) and FDPS (0.93, P = 1e-4), whereas the cholesterol-uptake "
    "receptor VLDLR (-0.52, P < 1e-4), PRKAA1 (-0.51, P = 0.026), LIMA1 "
    "(-0.22, P = 2.4e-3) and NSDHL (-0.19, P = 0.022) were down-regulated, "
    "indicating coordinated pathway remodelling during tumourigenesis "
    "rather than uniform up-regulation (Supplementary Figure S9; "
    "Supplementary Table S19). In a larger comparison of TCGA-BRCA "
    "tumours (n = 1,092) with GTEx normal breast tissue (n = 179), the "
    "up-regulation of DHCR24, DHCR7, FDPS and ABCG1 and the "
    "down-regulation of VLDLR, PRKAA1, LIMA1 and FDXR were confirmed "
    "with consistent directions in both cohorts (all Mann-Whitney "
    "P < 1e-12; G6PD and NSDHL were higher in tumours only in the Xena "
    "comparison, and HMGCS2 was directionally consistent but not "
    "significant; Supplementary Figure S11; Supplementary Table S21).")

REPL["Data availability."] = (
    "Data availability. ",
    "All data are public: TCGA-BRCA (NCI GDC; accessed 4-6 August 2026) "
    "and TCGA-CDR [14]; METABRIC (cBioPortal) [15]; PAM50 calls (UCSC "
    "Xena); FDPS HM450 methylation (cBioPortal); GSE21653 (NCBI GEO); "
    "GSE7390 (ArrayExpress E-GEOD-7390); GSE20711 (NCBI GEO); GSE15852 "
    "(NCBI GEO; paired normal-tumour); GSE91061 (NCBI GEO; immunotherapy "
    "RNA-seq); GSE176078 (CELLxGENE); GSE161529 (Mendeley Data mirror of "
    "the Pal et al. atlas [12]); TCGA+GTEx normal-tumour expression "
    "(UCSC Xena TcgaTargetGtex_rsem_gene_tpm); eQTLGen [35]; GTEx v8 "
    "[38,39]; BCAC GWAS (GWAS Catalog GCST010098/GCST010100/GCST004988; "
    "Zhang et al. 2020 [36] and Michailidou et al. 2017 [37]); 1000 "
    "Genomes Phase 3 EUR. A complete data inventory with accessions and "
    "usage notes is provided in Supplementary Table S17. Curated "
    "intermediate results are provided in the GitHub repository "
    "(https://github.com/y0307Yd/brca-cholesterol-signature) and the "
    "Zenodo archive (https://doi.org/10.5281/zenodo.21873168); raw data "
    "must be obtained from the original sources under their respective "
    "data-use policies. Supplementary Tables S1-S21 and Supplementary "
    "Figures S1-S11 accompany this manuscript; the TRIPOD-AI checklist "
    "(Supplementary Table S12), a comparison with published "
    "lipid-metabolism signatures (S11) and a graphical abstract are also "
    "provided in the repository.")

REPL["2.1 Data sources and verification"] = (
    "2.1 Data sources and verification ",
    "The overall study design and the flow of analyses are summarised in "
    "Supplementary Figure S2. For validation analyses, PAM50 calls for "
    "TCGA-BRCA were obtained from UCSC Xena (GDC hub) [9]; FDPS HM450 "
    "methylation beta values were obtained from the cBioPortal data "
    "repository (brca_tcga_methylation_hm450); the independent expression "
    "cohorts GSE21653 (Affymetrix HG-U133 Plus 2.0, NCBI GEO with GPL570 "
    "annotation) [10] and GSE7390 (Affymetrix HG-U133A, ArrayExpress "
    "E-GEOD-7390 processed matrix and SDRF clinical data, probes mapped "
    "with hgu133a.db) [11] were used; a fourth independent Affymetrix "
    "cohort (GSE20711, NCBI GEO) was used for classifier transfer; the "
    "paired normal-tumour cohort GSE15852 (Affymetrix HG-U133A; 43 pairs) "
    "was used for normal-versus-tumour expression comparisons; the "
    "immunotherapy cohort GSE91061 (RNA-seq, melanoma anti-PD-1; 51 "
    "pre-treatment samples) was used for exploratory immunotherapy-"
    "response associations; the UCSC Xena TCGA+GTEx TPM matrix "
    "(TcgaTargetGtex_rsem_gene_tpm) was used for the large-sample "
    "normal-tumour comparison (TCGA-BRCA n = 1,092; GTEx breast n = 179); "
    "and the second single-cell atlas GSE161529 (Mendeley Data mirror of "
    "the Pal et al. atlas) [12] was used for single-cell validation. The "
    "Affymetrix annotation was taken from the GEO GPL570 platform file. "
    "TCGA-BRCA RNA-sequencing STAR counts, methylation, mutation and "
    "clinical files were obtained from the NCI GDC portal [13] and "
    "verified against the official manifest by MD5 checksums (4,393/4,393 "
    "files); progression-free interval (PFI) was taken from TCGA-CDR "
    "[14]. METABRIC expression and clinical annotation with "
    "recurrence-free survival (RFS) were obtained from cBioPortal [15]. "
    "TCGA ER/PR/HER2 status was parsed from clinical XML; METABRIC "
    "receptor status from clinical sample files.")


def main():
    doc = docx.Document(SRC)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    for key, (bold, body) in REPL.items():
        hits = [p for p in paras if p.text.strip().startswith(key)]
        assert len(hits) == 1, f"{key!r}: {len(hits)} matches"
        set_text(hits[0], bold, body)
    for key, sentence in APPEND:
        hits = [p for p in paras if p.text.strip().startswith(key)]
        assert len(hits) == 1, f"append {key!r}: {len(hits)} matches"
        hits[0].add_run(" " + sentence)
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    for key, new_paras in INSERT_AFTER:
        anchor = next(p for p in paras if p.text.strip().startswith(key))
        for bold, body in new_paras:
            el = doc.element.body.makeelement(qn("w:p"), {})
            anchor._p.addnext(el)
            np_ = Paragraph(el, doc)
            set_text(np_, bold, body)
            anchor = np_
    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()
