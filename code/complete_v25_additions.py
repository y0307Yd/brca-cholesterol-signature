from __future__ import annotations

import re
import os
from pathlib import Path

import docx
os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parent / ".mplconfig"))
import matplotlib
import numpy as np
import pandas as pd
from scipy import stats
from sklearn.linear_model import LogisticRegression

matplotlib.use("Agg")
import matplotlib.pyplot as plt


REPO = Path(__file__).resolve().parents[1]
SOURCE = Path(os.environ.get("BRCA_CHOL_PROJECT_ROOT", REPO)).resolve()
OLD_OUT = SOURCE / "outputs" / "chol_metab_signature"
OUT = Path(os.environ.get("BRCA_CHOL_OUTPUT_DIR", REPO / "reproduction_output")).resolve()
OUT.mkdir(parents=True, exist_ok=True)


def bh_adjust(values: list[float]) -> np.ndarray:
    p = np.asarray(values, dtype=float)
    order = np.argsort(p)
    ranked = p[order]
    adj = ranked * len(p) / np.arange(1, len(p) + 1)
    adj = np.minimum.accumulate(adj[::-1])[::-1]
    result = np.empty_like(adj)
    result[order] = np.minimum(adj, 1.0)
    return result


def parse_estimate() -> pd.DataFrame:
    path = SOURCE / "work" / "estimate_scores.txt"
    tab = pd.read_csv(path, sep="\t", skiprows=2, index_col=0)
    tab = tab.drop(columns=[c for c in tab.columns if c == "Description"], errors="ignore")
    tab.columns = [str(c).replace(".", "-") for c in tab.columns]
    tab = tab.T
    tab.index = [re.sub(r"-(01[AB]).*$", r"-\1", x) for x in tab.index]
    return tab.apply(pd.to_numeric, errors="coerce")


def parse_pam50() -> pd.Series:
    path = SOURCE / "data" / "pam50" / "BRCA_clinicalMatrix.tsv"
    df = pd.read_csv(path, sep="\t", dtype=str, low_memory=False)
    s = df.set_index("sampleID")["PAM50Call_RNAseq"].dropna()
    s.index = ["-".join(str(x).split("-")[:3]) for x in s.index]
    return s[~s.index.duplicated(keep="first")]


def design_matrix(df: pd.DataFrame, include_subtype: bool, pam50: bool) -> np.ndarray:
    cols = [np.ones(len(df))]
    stage = pd.to_numeric(df["stage"], errors="coerce")
    cols.append(stage.to_numpy(float))
    cols.append(pd.to_numeric(df["ESTIMATEScore"], errors="coerce").to_numpy(float))
    if pam50:
        d = pd.get_dummies(df["PAM50"], drop_first=True, dtype=float)
        cols.extend([d[c].to_numpy() for c in d.columns])
    else:
        cols.append(pd.to_numeric(df["ER"], errors="coerce").to_numpy(float))
    if include_subtype:
        d = pd.get_dummies(df["subtype"].astype(str), drop_first=True, dtype=float)
        cols.extend([d[c].to_numpy() for c in d.columns])
    return np.column_stack(cols)


def partial_f(y: np.ndarray, x0: np.ndarray, x1: np.ndarray) -> tuple[float, float]:
    b0 = np.linalg.lstsq(x0, y, rcond=None)[0]
    b1 = np.linalg.lstsq(x1, y, rcond=None)[0]
    rss0 = float(np.sum((y - x0 @ b0) ** 2))
    rss1 = float(np.sum((y - x1 @ b1) ** 2))
    q = x1.shape[1] - x0.shape[1]
    df2 = len(y) - x1.shape[1]
    f = ((rss0 - rss1) / q) / (rss1 / df2)
    return f, float(stats.f.sf(f, q, df2))


def logistic_lrt(y: np.ndarray, x0: np.ndarray, x1: np.ndarray) -> float:
    m0 = LogisticRegression(penalty=None, max_iter=5000).fit(x0[:, 1:], y)
    m1 = LogisticRegression(penalty=None, max_iter=5000).fit(x1[:, 1:], y)
    p0 = np.clip(m0.predict_proba(x0[:, 1:])[:, 1], 1e-12, 1 - 1e-12)
    p1 = np.clip(m1.predict_proba(x1[:, 1:])[:, 1], 1e-12, 1 - 1e-12)
    ll0 = np.sum(y * np.log(p0) + (1 - y) * np.log(1 - p0))
    ll1 = np.sum(y * np.log(p1) + (1 - y) * np.log(1 - p1))
    return float(stats.chi2.sf(2 * (ll1 - ll0), x1.shape[1] - x0.shape[1]))


def analyse_tide() -> tuple[pd.DataFrame, pd.DataFrame, dict[str, float]]:
    tide = pd.read_csv(OLD_OUT / "tide_by_sample.csv", index_col=0)
    sub = pd.read_csv(OLD_OUT / "tcga_subtypes.csv")
    # tidepy may reorder columns. S<n> is the original zero-based expression column.
    source_row = tide.index.to_series().str.extract(r"^S(\d+)$", expand=False).astype(int)
    if source_row.duplicated().any() or source_row.min() != 0 or source_row.max() != len(sub) - 1:
        raise RuntimeError("TIDE synthetic sample identifiers do not map one-to-one to TCGA rows")
    tide["source_row"] = source_row.to_numpy()
    clinical = sub.reset_index(names="source_row")
    df = tide.merge(clinical, on="source_row", validate="one_to_one")

    est = parse_estimate().reset_index(names="estimate_id")
    est["sample_key"] = est["estimate_id"].str.replace("-", ".", regex=False)
    df["sample_key"] = df["sample_id"].str.replace("-", ".", regex=False)
    df = df.merge(est.drop(columns="estimate_id"), on="sample_key", how="left", validate="one_to_one")

    pam = parse_pam50()
    df["patient_key"] = df["sample_id"].str.split("-").str[:3].str.join("-")
    df["PAM50"] = df["patient_key"].map(pam)

    features = ["TIDE", "IFNG", "Dysfunction", "Exclusion", "MDSC", "CAF",
                "TAM M2", "CD274", "CD8", "CTL"]
    rows = []
    for feature in features:
        groups = [df.loc[df.subtype == k, feature].dropna().to_numpy() for k in range(1, 5)]
        kw_p = float(stats.kruskal(*groups).pvalue)
        row = {"feature": feature, "n": int(df[feature].notna().sum()), "unadjusted_P": kw_p}
        row.update({f"median_C{k}": float(np.median(g)) for k, g in zip(range(1, 5), groups)})

        cov = df[[feature, "subtype", "ER", "stage", "ESTIMATEScore"]].dropna()
        row["adjusted_n_ER_stage_ESTIMATE"] = len(cov)
        _, row["adjusted_P_ER_stage_ESTIMATE"] = partial_f(
            cov[feature].to_numpy(float), design_matrix(cov, False, False), design_matrix(cov, True, False))

        covp = df[[feature, "subtype", "stage", "ESTIMATEScore", "PAM50"]].dropna()
        row["adjusted_n_PAM50_stage_ESTIMATE"] = len(covp)
        _, row["adjusted_P_PAM50_stage_ESTIMATE"] = partial_f(
            covp[feature].to_numpy(float), design_matrix(covp, False, True), design_matrix(covp, True, True))
        rows.append(row)

    summary = pd.DataFrame(rows)
    fdr_names = {
        "unadjusted_P": "unadjusted_FDR",
        "adjusted_P_ER_stage_ESTIMATE": "adjusted_FDR_ER_stage_ESTIMATE",
        "adjusted_P_PAM50_stage_ESTIMATE": "adjusted_FDR_PAM50_stage_ESTIMATE",
    }
    for col, fdr_col in fdr_names.items():
        summary[fdr_col] = bh_adjust(summary[col].tolist())

    response = []
    for k, g in df.groupby("subtype"):
        response.append({"subtype": f"C{k}", "n": len(g), "predicted_responder_n": int(g.Responder.sum()),
                         "predicted_responder_pct": float(g.Responder.mean() * 100)})
    response = pd.DataFrame(response)
    ct = pd.crosstab(df.subtype, df.Responder)
    chi_p = float(stats.chi2_contingency(ct).pvalue)
    cov = df[["Responder", "subtype", "ER", "stage", "ESTIMATEScore"]].dropna()
    lrt_p = logistic_lrt(cov.Responder.astype(int).to_numpy(), design_matrix(cov, False, False),
                         design_matrix(cov, True, False))
    covp = df[["Responder", "subtype", "stage", "ESTIMATEScore", "PAM50"]].dropna()
    lrt_p_pam = logistic_lrt(covp.Responder.astype(int).to_numpy(), design_matrix(covp, False, True),
                             design_matrix(covp, True, True))

    summary.to_csv(OUT / "Supplementary_Table_S22_TIDE_subtype_sensitivity.csv", index=False,
                   encoding="utf-8-sig")
    response.to_csv(OUT / "TIDE_predicted_response_by_subtype.csv", index=False, encoding="utf-8-sig")
    df.to_csv(OUT / "TIDE_by_sample_corrected_mapping.csv", index=False, encoding="utf-8-sig")

    fig, axes = plt.subplots(2, 3, figsize=(11.5, 7.5))
    palette = ["#b23a48", "#277da1", "#43aa8b", "#f8961e"]
    for ax, feature in zip(axes.ravel(), ["TIDE", "IFNG", "Dysfunction", "Exclusion", "CD274", "CD8"]):
        data = [df.loc[df.subtype == k, feature].dropna() for k in range(1, 5)]
        bp = ax.boxplot(data, tick_labels=[f"C{k}" for k in range(1, 5)], widths=0.55,
                        patch_artist=True, showfliers=False)
        for box, color in zip(bp["boxes"], palette):
            box.set_facecolor(color); box.set_alpha(0.72)
        p = summary.loc[summary.feature == feature, "unadjusted_P"].iloc[0]
        p2 = summary.loc[summary.feature == feature, "adjusted_P_ER_stage_ESTIMATE"].iloc[0]
        ax.set_title(f"{feature}\nKW P={p:.2g}; adjusted P={p2:.2g}", fontsize=9)
        ax.grid(axis="y", alpha=0.2)
    fig.suptitle("TIDE-derived immune-evasion features by molecular subtype\n"
                 "Adjusted model: ER status, stage and ESTIMATE score", fontsize=12)
    fig.tight_layout(rect=(0, 0, 1, 0.94))
    fig.savefig(OUT / "Supplementary_Figure_S8_TIDE_subtypes_corrected.png", dpi=300,
                bbox_inches="tight")
    plt.close(fig)
    return summary, response, {"chi_p": chi_p, "lrt_p": lrt_p, "lrt_p_pam": lrt_p_pam,
                               "n_pam50": int(df.PAM50.notna().sum())}


def replace_paragraph(p, text: str) -> None:
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    p.add_run(text)


def patch_docx(summary: pd.DataFrame, response: pd.DataFrame, stats_: dict[str, float]) -> Path:
    src = OLD_OUT / "Manuscript_Cholesterol_Metabolism_BRCA_v25.docx"
    out = OUT / "Manuscript_Cholesterol_Metabolism_BRCA_v26.docx"
    doc = docx.Document(src)
    paras = [p for p in doc.paragraphs if p.text.strip()]

    duplicate = next(
        p for p in paras
        if p.text.strip().startswith("Expression was transformed as log2(CPM+1). The 8,000 most variable genes entered")
    )
    duplicate._element.getparent().remove(duplicate._element)

    p26 = next(p for p in paras if p.text.strip().startswith("2.6 Molecular subtyping"))
    p26.add_run(
        " As an exploratory sensitivity analysis, TIDE-derived immune-evasion features and predicted "
        "checkpoint-blockade response were compared across subtypes. Because TIDE reorders its synthetic "
        "sample identifiers, outputs were remapped explicitly by the zero-based S<n> identifier before "
        "merging. Continuous features were tested by Kruskal-Wallis and by nested linear models adjusted "
        "for ER status, stage and ESTIMATE score; a second model replaced ER status with PAM50. Predicted "
        "response was analysed analogously using likelihood-ratio tests. These predictions were not treated "
        "as observed treatment outcomes."
    )

    p35 = next(p for p in paras if p.text.strip().startswith("3.5 Immune microenvironment"))
    tide = summary.loc[summary.feature == "TIDE"].iloc[0]
    significant_pam = summary.loc[
        summary.adjusted_FDR_PAM50_stage_ESTIMATE < 0.05, "feature"
    ].tolist()
    rates = ", ".join(f"{r.subtype} {r.predicted_responder_pct:.1f}%" for r in response.itertuples())
    p35.add_run(
        f" In the corrected sample-level TIDE sensitivity analysis, the composite TIDE score differed "
        f"across subtypes (Kruskal-Wallis P = {tide.unadjusted_P:.2e}) and was lowest in C1 "
        f"(median {tide.median_C1:.2f}), consistent with lower predicted immune escape. The association "
        f"persisted after adjustment for ER status, stage and ESTIMATE score "
        f"(P = {tide.adjusted_P_ER_stage_ESTIMATE:.2e}) and after replacing ER with PAM50 "
        f"(P = {tide.adjusted_P_PAM50_stage_ESTIMATE:.2e}). In the PAM50-adjusted model, "
        f"{', '.join(significant_pam)} remained significant at FDR < 0.05. Binary TIDE response calls "
        f"were nearly invariant ({rates}; unadjusted chi-square P = {stats_['chi_p']:.3f}; adjusted "
        f"likelihood-ratio P = {stats_['lrt_p']:.3f}) and were therefore not clinically informative. "
        f"These findings support subtype differences in predicted immune-evasion biology, not observed "
        f"checkpoint-blockade benefit (Supplementary Figure S8; Supplementary Table S22)."
    )

    p316 = next(p for p in paras if p.text.strip().startswith("3.16 Cholesterol genes are remodelled"))
    corrected = (
        "3.16 Cholesterol genes are remodelled in tumour versus normal breast tissue\n"
        "In a paired normal-tumour cohort (GSE15852, 43 pairs), DHCR24 (log2 fold change 0.59, paired "
        "Wilcoxon P = 7e-4), DHCR7 (0.82, P = 2.8e-3), HSD17B7 (0.36, P = 1.2e-3) and FDPS "
        "(0.93, P = 1e-4) were significantly up-regulated in tumours; HMGCS2 showed the same direction "
        "but was not significant (0.83, P = 0.13). VLDLR (-0.52, P < 1e-4), PRKAA1 (-0.51, "
        "P = 0.026), LIMA1 (-0.22, P = 2.4e-3) and NSDHL (-0.19, P = 0.022) were down-regulated, "
        "indicating pathway remodelling rather than uniform up-regulation (Supplementary Figure S9; "
        "Supplementary Table S19). In the larger TCGA-BRCA (n = 1,092) versus GTEx normal breast "
        "comparison (n = 179), DHCR24, DHCR7 and FDPS were again higher, whereas VLDLR, PRKAA1 and "
        "LIMA1 were again lower (all Mann-Whitney P < 1e-12). ABCG1 and FDXR were significant only in "
        "the Xena comparison; G6PD and NSDHL also differed there, while HMGCS2 remained non-significant "
        "(Supplementary Figure S11; Supplementary Table S21)."
    )
    replace_paragraph(p316, corrected)

    figure_refs = {
        "3.2 WGCNA and the null module-by-pathway intersection": " The network diagnostics and module-trait associations are shown in Figures 1 and 2.",
        "3.3 Hub genes and the ER-status classifier": " Model coefficients, survival-stratified curves, external discrimination and ER-score separation are shown in Figures 4 and 5.",
        "3.4 Molecular subtypes defined by hub genes": " Consensus diagnostics and subtype profiles are shown in Figure 6.",
        "3.12 PAM50 calls recapitulate the hub-gene molecular subtypes": " PAM50 composition is shown in Figure 11.",
        "3.13 Independent validation in GSE21653": " The GSE7390 and GSE21653 validation results are shown in Figures 12 and 13, respectively.",
        "3.15 FDPS DNA methylation is inversely associated": " The methylation-expression analyses are shown in Figures 15 and 16.",
    }
    for key, sentence in figure_refs.items():
        next(p for p in paras if p.text.strip().startswith(key)).add_run(sentence)

    p43 = next(p for p in paras if p.text.strip().startswith("4.3 Subtypes, proliferation and immunity"))
    p43.add_run(
        " TIDE-based sensitivity analyses supported lower predicted immune escape in C1 after adjustment "
        "for ER or PAM50, stage and ESTIMATE score. However, the nearly invariant binary response calls "
        "and absence of treated breast-cancer outcomes preclude inference of differential clinical benefit."
    )

    pav = next(p for p in paras if p.text.strip().startswith("Data availability."))
    replace_paragraph(pav, pav.text.replace("Supplementary Tables S1-S21", "Supplementary Tables S1-S22"))
    doc.save(out)
    return out


def main() -> None:
    summary, response, stats_ = analyse_tide()
    manuscript = patch_docx(summary, response, stats_)
    report = OUT / "v26_completion_report.md"
    report.write_text(
        "# v26 completion report\n\n"
        "- Corrected the TIDE sample mapping defect in the previous exploratory script.\n"
        "- Added unadjusted, ER/stage/ESTIMATE-adjusted and PAM50/stage/ESTIMATE-adjusted analyses.\n"
        "- Added Supplementary Figure S8 and Supplementary Table S22.\n"
        "- Corrected HMGCS2 and cross-cohort normal-versus-tumour wording in section 3.16.\n"
        "- Added explicit limitations preventing TIDE predictions from being described as clinical response.\n"
        "- GDSC was not added because no auditable local drug-response data were available.\n\n"
        f"PAM50 available for {stats_['n_pam50']} TIDE-mapped samples.\n",
        encoding="utf-8",
    )
    print(manuscript)
    print(summary.to_string(index=False))
    print(response.to_string(index=False))
    print(stats_)


if __name__ == "__main__":
    main()
