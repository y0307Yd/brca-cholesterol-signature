# -*- coding: utf-8 -*-
"""Generate v23 from v22: pathway-activity panel + Discussion correction."""
import docx

SRC = (r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature"
       r"\Manuscript_Cholesterol_Metabolism_BRCA_v22.docx")
OUT = (r"C:\Users\Y\Documents\Codex\2026-08-06\new-chat\outputs\chol_metab_signature"
       r"\Manuscript_Cholesterol_Metabolism_BRCA_v23.docx")


def set_text(p, bold_part, body_part):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if bold_part:
        r = p.add_run(bold_part.rstrip("\n"))
        r.bold = True
        if bold_part.endswith("\n"):
            p.add_run().add_break()
    lines = body_part.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        if line:
            p.add_run(line)


REPL = {
    "The four cholesterol-metabolism subtypes recapitulate": (
        "",
        "The four cholesterol-metabolism subtypes recapitulate known "
        "breast cancer biology: C1 is predominantly ER-negative, "
        "proliferative (cell-cycle up) and immune-infiltrated (highest "
        "CD8 T, regulatory T, NK, B-cell and macrophage scores), "
        "consistent with the immunologically hot, proliferative phenotype "
        "of triple-negative tumours; C3 is almost uniformly ER+, "
        "immune-cold and oxidative-phosphorylation-high; C2 shows "
        "down-regulated cell-cycle programs and high stromal/dendritic "
        "scores; C4 is a mixed, neutrophil-high group. Rank-based "
        "pathway-activity scores showed the strongest subtype differences "
        "for estrogen response (highest in the luminal-like C2/C3) and "
        "for cell-cycle and immune-cytokine programs (highest in C1), "
        "whereas cholesterol biosynthesis was only modestly higher in C4 "
        "than in C1 (Supplementary Figure S7). These patterns provide a "
        "cell-level interpretation of the signature even before "
        "single-cell data are added: the cholesterol-synthesis-low, "
        "ER-low C1 subtype is proliferative and immune-active [2,8], "
        "while cholesterol-synthesis activity per se does not segregate "
        "the luminal subtypes, indicating that the ER-classifier signal "
        "reflects ER-pathway biology more than cholesterol-pathway "
        "activity as a whole."),
    "Data availability.": (
        "Data availability. ",
        "All data are public: TCGA-BRCA (NCI GDC; accessed 4-6 August "
        "2026) and TCGA-CDR [14]; METABRIC (cBioPortal) [15]; PAM50 calls "
        "(UCSC Xena); FDPS HM450 methylation (cBioPortal); GSE21653 (NCBI "
        "GEO); GSE7390 (ArrayExpress E-GEOD-7390); GSE20711 (NCBI GEO); "
        "GSE176078 (CELLxGENE); GSE161529 (Mendeley Data mirror of the "
        "Pal et al. atlas [12]); eQTLGen [35]; GTEx v8 [38,39]; BCAC GWAS "
        "(GWAS Catalog GCST010098/GCST010100/GCST004988; Zhang et al. "
        "2020 [36] and Michailidou et al. 2017 [37]); 1000 Genomes Phase "
        "3 EUR. A complete data inventory with accessions and usage notes "
        "is provided in Supplementary Table S17. Curated intermediate "
        "results are provided in the GitHub repository "
        "(https://github.com/y0307Yd/brca-cholesterol-signature) and the "
        "Zenodo archive (https://doi.org/10.5281/zenodo.21873168); raw "
        "data must be obtained from the original sources under their "
        "respective data-use policies. Supplementary Tables S1-S18 and "
        "Supplementary Figures S1-S7 accompany this manuscript; the "
        "TRIPOD-AI checklist (Supplementary Table S12), a comparison with "
        "published lipid-metabolism signatures (S11) and a graphical "
        "abstract are also provided in the repository."),
}

APPEND = [
    ("3.5 Immune microenvironment and pathway programs",
     "Rank-based pathway-activity scores (ssGSEA, alpha = 0.25) showed "
     "the largest subtype differences for estrogen response (highest in "
     "C2/C3), cell cycle and immune cytokines (highest in C1), with "
     "smaller differences for cholesterol biosynthesis (highest in C4, "
     "lowest in C1) and fatty-acid oxidation (all Kruskal-Wallis "
     "P < 0.001; Supplementary Figure S7)."),
]


def main():
    doc = docx.Document(SRC)
    paras = [p for p in doc.paragraphs if p.text.strip()]
    for key, (bold, body) in REPL.items():
        hits = [p for p in paras if p.text.strip().startswith(key)]
        assert len(hits) == 1, f"{key!r}: {len(hits)} matches"
        set_text(hits[0], bold, body)
    for key, sentence in APPEND:
        hits = [p for p in paras if p.text.strip().startswith(key)]
        assert len(hits) == 1, f"append {key!r}: {len(hits)} matches"
        hits[0].add_run(" " + sentence)
    doc.save(OUT)
    print("saved", OUT)


if __name__ == "__main__":
    main()
